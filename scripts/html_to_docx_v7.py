"""Convert the restructured editorial HTML v7 to DOCX."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

INPUT = r"C:\Users\INTEL\Desktop\fomento-audiovisual\output_final\Uma política de fomento baseada em evidências.html"
OUTPUT = r"C:\Users\INTEL\Desktop\fomento-audiovisual\output_final\Uma política de fomento baseada em evidências.docx"

with open(INPUT, "r", encoding="utf-8") as f:
    html = f.read()

# Extract article
art_start = html.find("<article>")
art_end = html.find("</article>")
article = html[art_start:art_end]

# Remove embedded images and charts
article = re.sub(r'<img[^>]*src="data:image[^"]*"[^>]*>', '[grafico - ver versao HTML]', article)
article = re.sub(r'<!-- chart\d+-plotly -->.*?(?=</p>)', '[grafico interativo - ver versao HTML]', article, flags=re.DOTALL)
# Remove Plotly script calls
article = re.sub(r'<p>.*?Plotly\.newPlot.*?</p>', '<p>[grafico interativo - ver versao HTML]</p>', article, flags=re.DOTALL)

doc = Document()

# ── Styles ─────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

doc.styles["Heading 1"].font.size = Pt(24)
doc.styles["Heading 2"].font.size = Pt(16)
doc.styles["Heading 3"].font.size = Pt(13)
doc.styles["Heading 4"].font.size = Pt(11)


def strip_tags(s):
    s = re.sub(r"<br\s*/?>", "\n", s)
    s = re.sub(r"<[^>]+>", "", s)
    s = s.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    s = s.replace("&nbsp;", " ")
    return s.strip()


def add_rich_para(doc, html_text, style_name="Normal"):
    p = doc.add_paragraph(style=style_name)
    parts = re.split(r"(<strong[^>]*>|</strong>|<em[^>]*>|</em>|<a [^>]*>|</a>)", html_text)
    is_bold = False
    is_italic = False
    for part in parts:
        if part.startswith("<strong"):
            is_bold = True
            continue
        if part == "</strong>":
            is_bold = False
            continue
        if part.startswith("<em"):
            is_italic = True
            continue
        if part == "</em>":
            is_italic = False
            continue
        if part.startswith("<a ") or part == "</a>":
            continue
        text = strip_tags(part)
        if text:
            run = p.add_run(text)
            run.bold = is_bold
            run.italic = is_italic
    return p


def add_section_label(doc, text, color=RGBColor(0x6C, 0x7B, 0xF7)):
    """Add a colored label paragraph (like 'EVIDENCIA 1' or 'ARGUMENTO')."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_evidence_box(doc, ev_items, lim_items, conclusion):
    """Add a formatted evidence summary as a table."""
    # Use a 3-column table
    t = doc.add_table(rows=2, cols=3)
    t.style = "Table Grid"

    # Headers
    headers = ["EVIDENCIA", "LIMITACOES", "CONCLUSAO"]
    colors = [RGBColor(0x6C, 0x7B, 0xF7), RGBColor(0xC8, 0x9A, 0x6A), RGBColor(0x5A, 0x9A, 0x6A)]
    for ci, (header, color) in enumerate(zip(headers, colors)):
        cell = t.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(8)
        run.bold = True
        run.font.color.rgb = color
        run.font.name = "Calibri"

    # Content
    ev_text = "\n".join(f"- {strip_tags(e)}" for e in ev_items)
    lim_text = "\n".join(f"- {strip_tags(l)}" for l in lim_items)
    conc_text = strip_tags(conclusion)

    for ci, text in enumerate([ev_text, lim_text, conc_text]):
        cell = t.rows[1].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Georgia"

    doc.add_paragraph("")  # spacing
    return t


# ── Parse and build DOCX ──────────────────────────────────────────
lines = article.split("\n")
in_blockquote = False
in_table = False
table_rows = []
in_evidence_box = False
eb_col_idx = 0
eb_items = [[], [], []]
in_collapsible = False
claim_count = 0

i = 0
while i < len(lines):
    s = lines[i].strip()

    if not s:
        i += 1
        continue

    # Skip structural tags
    if s.startswith("<article") or s.startswith("</article"):
        i += 1
        continue
    if s == "</div>" or s == "<div>":
        i += 1
        continue

    # Skip chart containers
    if 'style="display:block;width:calc' in s:
        i += 1
        continue

    # HTML comments (section markers)
    if s.startswith("<!--"):
        i += 1
        continue

    # ── Structural wrappers ──
    # Claim section
    if 'class="claim-section"' in s:
        i += 1
        continue

    # Claim header
    if 'class="claim-header"' in s:
        i += 1
        continue

    if 'class="claim-num"' in s:
        num = strip_tags(s)
        claim_count = int(num) if num.isdigit() else claim_count + 1
        i += 1
        continue

    # Context section
    if 'class="context-section"' in s:
        add_section_label(doc, "CONTEXTO", RGBColor(0xC8, 0x9A, 0x6A))
        i += 1
        continue

    # Opinion section
    if 'class="opinion-section"' in s:
        i += 1
        continue
    if 'class="opinion-label"' in s:
        text = strip_tags(s)
        add_section_label(doc, text.upper(), RGBColor(0xC8, 0x9A, 0x6A))
        i += 1
        continue

    # Aside note
    if 'class="aside-note"' in s:
        i += 1
        continue

    # Collapsible
    if 'class="collapsible"' in s:
        in_collapsible = True
        i += 1
        continue
    if in_collapsible and "</details>" in s:
        in_collapsible = False
        i += 1
        continue
    if "<summary>" in s or "</summary>" in s:
        text = strip_tags(s)
        if text:
            doc.add_heading(text, level=2)
        i += 1
        continue
    if 'class="collapsible-body"' in s:
        i += 1
        continue

    # Evidence box
    if 'class="evidence-box"' in s:
        in_evidence_box = True
        eb_items = [[], [], []]
        eb_col_idx = 0
        i += 1
        continue
    if in_evidence_box:
        if 'class="eb-col"' in s:
            i += 1
            continue
        if 'class="eb-label"' in s:
            i += 1
            continue
        if 'class="eb-grid"' in s:
            i += 1
            continue
        if s == "</div>":
            # Could be end of eb-col or end of evidence-box
            # Count divs
            i += 1
            continue
        if s.startswith("<li"):
            text = strip_tags(s)
            eb_items[eb_col_idx].append(text)
            i += 1
            continue
        if s.startswith("<p") and 'class="eb' not in s:
            text = strip_tags(s)
            if text:
                eb_items[eb_col_idx].append(text)
            i += 1
            continue
        if s in ("<ul>", "</ul>"):
            if s == "</ul>" and eb_col_idx < 2:
                eb_col_idx += 1
            i += 1
            continue
        # End of evidence box - detect by next significant element
        if s.startswith("<p>") or s.startswith("<h") or s.startswith("<div class"):
            in_evidence_box = False
            add_evidence_box(doc, eb_items[0], eb_items[1],
                           eb_items[2][0] if eb_items[2] else "")
            continue  # don't increment, reprocess this line
        i += 1
        continue

    # ── Hero ──
    if 'class="article-hero"' in s:
        i += 1
        continue
    if 'class="epigraph"' in s:
        i += 1
        continue
    if "proposals-grid" in s or "proposal-card" in s or "pc-num" in s or "pc-body" in s:
        i += 1
        continue

    # ── Sobre este texto ──
    if 'id="sobre-este-texto"' in s:
        p = doc.add_paragraph("")
        run = p.add_run("SOBRE ESTE TEXTO")
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        i += 1
        continue
    if "Sobre este texto" in s and "font-size:11px" in s:
        i += 1
        continue

    # ── Headings ──
    m = re.match(r"<h1[^>]*>(.*?)</h1>", s)
    if m:
        doc.add_heading(strip_tags(m.group(1)), level=0)
        i += 1
        continue

    if "article-meta" in s:
        text = strip_tags(s)
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(10)
        i += 1
        continue

    m = re.match(r"<h2[^>]*>(.*?)</h2>", s)
    if m:
        title = strip_tags(m.group(1))
        doc.add_heading(title, level=1)
        i += 1
        continue

    m = re.match(r"<h3[^>]*>(.*?)</h3>", s)
    if m:
        doc.add_heading(strip_tags(m.group(1)), level=2)
        i += 1
        continue

    m = re.match(r"<h4[^>]*>(.*?)</h4>", s)
    if m:
        doc.add_heading(strip_tags(m.group(1)), level=3)
        i += 1
        continue

    # ── Blockquote ──
    if "<blockquote" in s:
        in_blockquote = True
        i += 1
        continue
    if "</blockquote>" in s:
        in_blockquote = False
        i += 1
        continue

    if s.startswith("<cite"):
        text = strip_tags(s)
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Cm(1.5)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        i += 1
        continue

    # ── Table ──
    if "<table" in s:
        in_table = True
        table_rows = []
        i += 1
        continue
    if "</table>" in s:
        in_table = False
        if table_rows:
            ncols = max(len(r) for r in table_rows)
            t = doc.add_table(rows=len(table_rows), cols=ncols)
            t.style = "Table Grid"
            for ri, row_data in enumerate(table_rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < ncols:
                        t.rows[ri].cells[ci].text = cell_text
            doc.add_paragraph("")
        i += 1
        continue
    if in_table:
        if "<tr" in s:
            table_rows.append([])
        elif "<th" in s or "<td" in s:
            text = strip_tags(s)
            if table_rows:
                table_rows[-1].append(text)
        i += 1
        continue

    # ── Lists ──
    if s.startswith("<ul") or s.startswith("<ol"):
        i += 1
        continue
    if s in ("</ul>", "</ol>"):
        i += 1
        continue
    if s.startswith("<li"):
        text = strip_tags(s)
        if text:
            add_rich_para(doc, s, style_name="List Bullet")
        i += 1
        continue

    # ── HR ──
    if s.startswith("<hr"):
        p = doc.add_paragraph("_" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
        i += 1
        continue

    # ── Paragraphs ──
    if s.startswith("<p"):
        text = strip_tags(s)
        if not text or len(text) < 3:
            i += 1
            continue
        if "grafico" in text and "ver versao HTML" in text:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.italic = True
                run.font.size = Pt(9)
            i += 1
            continue
        if in_blockquote:
            p = add_rich_para(doc, s)
            p.paragraph_format.left_indent = Cm(1.5)
            for run in p.runs:
                run.italic = True
            i += 1
            continue
        if "font-size:0.88em" in lines[i] or (text.startswith("Nota:") and len(text) < 500):
            p = add_rich_para(doc, s)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue
        add_rich_para(doc, s)
        i += 1
        continue

    # ── Remaining text ──
    text = strip_tags(s)
    if text and len(text) > 10 and not s.startswith("<"):
        add_rich_para(doc, s)

    i += 1

doc.save(OUTPUT)
print(f"DOCX saved: {OUTPUT}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
