"""Convert the editorial HTML to DOCX for editing."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT = r"C:\Users\INTEL\Desktop\fomento-audiovisual\output_final\Uma política de fomento baseada em evidências_v6_editado_editorial.html"
OUTPUT = r"C:\Users\INTEL\Desktop\fomento-audiovisual\output_final\Uma política de fomento baseada em evidências_v6.docx"

with open(INPUT, "r", encoding="utf-8") as f:
    html = f.read()

# Extract article content
art_start = html.find("<article>")
art_end = html.find("</article>")
article = html[art_start:art_end]

# Remove embedded images (base64)
article = re.sub(r'<img[^>]*src="data:image[^"]*"[^>]*>', '[grafico - ver versao HTML]', article)
# Remove Plotly chart divs
article = re.sub(r'<!-- chart\d+-plotly -->.*?(?=</p>)', '[grafico interativo - ver versao HTML]', article, flags=re.DOTALL)

doc = Document()

# Set default font
style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

doc.styles["Heading 1"].font.size = Pt(22)
doc.styles["Heading 2"].font.size = Pt(16)
doc.styles["Heading 3"].font.size = Pt(13)


def strip_tags(s):
    s = re.sub(r"<br\s*/?>", "\n", s)
    s = re.sub(r"<[^>]+>", "", s)
    s = s.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    s = s.replace("&nbsp;", " ")
    return s.strip()


def add_rich_para(doc, html_text, style_name="Normal"):
    p = doc.add_paragraph(style=style_name)
    parts = re.split(r"(<strong[^>]*>|</strong>|<em[^>]*>|</em>|<a [^>]*>|</a>)", html_text)
    is_bold = False
    is_italic = False
    for part in parts:
        if part.startswith("<strong"):
            is_bold = True
            continue
        if part == "</strong>":
            is_bold = False
            continue
        if part.startswith("<em"):
            is_italic = True
            continue
        if part == "</em>":
            is_italic = False
            continue
        if part.startswith("<a "):
            continue
        if part == "</a>":
            continue
        text = strip_tags(part)
        if text:
            run = p.add_run(text)
            run.bold = is_bold
            run.italic = is_italic
    return p


lines = article.split("\n")
in_blockquote = False
in_table = False
table_rows = []
skip_div = 0

for line in lines:
    s = line.strip()
    if not s:
        continue

    if s.startswith("<article") or s.startswith("</article"):
        continue
    if s in ("</div>", "<div>"):
        if skip_div > 0:
            skip_div -= 1
        continue

    # Skip chart container divs
    if 'style="display:block;width:calc' in s:
        continue

    # Article hero
    if 'class="article-hero"' in s:
        continue
    if 'class="epigraph"' in s:
        continue

    # Sobre este texto box
    if 'id="sobre-este-texto"' in s:
        p = doc.add_paragraph("")
        run = p.add_run("SOBRE ESTE TEXTO")
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        continue

    # Sobre este texto label
    if "Sobre este texto" in s and "font-size:11px" in s:
        continue

    # Proposal cards
    if "proposal-card" in s or "pc-num" in s or "pc-body" in s:
        continue
    if "proposals-grid" in s:
        continue

    # Title
    m = re.match(r"<h1[^>]*>(.*?)</h1>", s)
    if m:
        doc.add_heading(strip_tags(m.group(1)), level=0)
        continue

    # Meta
    if "article-meta" in s:
        text = strip_tags(s)
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(10)
        continue

    # H2
    m = re.match(r"<h2[^>]*>(.*?)</h2>", s)
    if m:
        doc.add_heading(strip_tags(m.group(1)), level=1)
        continue

    # H3
    m = re.match(r"<h3[^>]*>(.*?)</h3>", s)
    if m:
        doc.add_heading(strip_tags(m.group(1)), level=2)
        continue

    # H4
    m = re.match(r"<h4[^>]*>(.*?)</h4>", s)
    if m:
        doc.add_heading(strip_tags(m.group(1)), level=3)
        continue

    # Blockquote
    if "<blockquote" in s:
        in_blockquote = True
        continue
    if "</blockquote>" in s:
        in_blockquote = False
        continue

    # Cite
    if s.startswith("<cite"):
        text = strip_tags(s)
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Cm(1.5)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        continue

    # Table
    if "<table" in s:
        in_table = True
        table_rows = []
        continue
    if "</table>" in s:
        in_table = False
        if table_rows:
            ncols = max(len(r) for r in table_rows)
            t = doc.add_table(rows=len(table_rows), cols=ncols)
            t.style = "Table Grid"
            for ri, row_data in enumerate(table_rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < ncols:
                        t.rows[ri].cells[ci].text = cell_text
            doc.add_paragraph("")
        continue
    if in_table:
        if "<tr" in s:
            table_rows.append([])
        elif "<th" in s or "<td" in s:
            text = strip_tags(s)
            if table_rows:
                table_rows[-1].append(text)
        continue

    # Lists
    if s.startswith("<ul") or s.startswith("<ol"):
        continue
    if s in ("</ul>", "</ol>"):
        continue
    if s.startswith("<li"):
        text = strip_tags(s)
        if text:
            add_rich_para(doc, s, style_name="List Bullet")
        continue

    # HR
    if s.startswith("<hr"):
        p = doc.add_paragraph("_" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
        continue

    # Paragraphs
    if s.startswith("<p"):
        text = strip_tags(s)
        if not text or len(text) < 3:
            continue
        # Chart placeholder
        if "grafico" in text and "ver versao HTML" in text:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.italic = True
                run.font.size = Pt(9)
            continue
        # Blockquote paragraph
        if in_blockquote:
            p = add_rich_para(doc, s)
            p.paragraph_format.left_indent = Cm(1.5)
            for run in p.runs:
                run.italic = True
            continue
        # Note paragraphs
        if "font-size:0.88em" in line or (text.startswith("Nota:") and len(text) < 500):
            p = add_rich_para(doc, s)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue
        # Normal paragraph
        add_rich_para(doc, s)
        continue

    # Bullet-style paragraphs
    text = strip_tags(s)
    if text and len(text) > 10:
        # Catch any remaining text content
        if not s.startswith("<"):
            add_rich_para(doc, s)

doc.save(OUTPUT)
print(f"DOCX saved: {OUTPUT}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
