from __future__ import annotations

import math
import json
import re
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DOC = ROOT / "output_final" / "Análise de Dados - Fomento Audiovisual Brasileiro.docx"
CHART_DIR = ROOT / "resultados" / "doc_analise_graficos"
CHART_DIR.mkdir(parents=True, exist_ok=True)


COLORS = {
    "blue": "#2f5f8f",
    "orange": "#c7792b",
    "green": "#2d7f5e",
    "red": "#b34a4a",
    "purple": "#6f5aa7",
    "gray": "#6f7782",
    "yellow": "#c5a12d",
    "teal": "#258a8a",
}

CLUSTER_COLORS = {
    "Duplo Retorno": COLORS["green"],
    "Retorno Doméstico": COLORS["orange"],
    "Retorno Internacional": COLORS["blue"],
    "Fomento Baixo Retorno": COLORS["red"],
    "Pequeno Porte": COLORS["gray"],
}


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(ROOT / "resultados" / "datasets" / name, sep=";", encoding="utf-8-sig")


obra = read_csv("base_nivel_obra.csv")
prod = read_csv("base_nivel_produtora.csv")
cham = read_csv("base_nivel_chamada.csv")
inv = read_csv("base_nivel_investimento.csv")
master = pd.read_excel(ROOT / "resultados" / "tabela_consolidada_obras.xlsx", sheet_name="Obras")
direcao = read_csv("base_nivel_direcao.csv")
curtas = pd.read_excel(ROOT / "dados" / "curtas_brasileiros_festivais_internacionais.xlsx", sheet_name="Dados")
citacoes = pd.read_csv(ROOT / "dados" / "citacoes_diretores.csv", encoding="utf-8-sig")

for df in [obra, prod, cham, inv, master]:
    for col in df.columns:
        if df[col].dtype == object:
            continue

for col in [
    "investimento_fsa_deflac",
    "investimento_renuncia_total_deflac",
    "investimento_total_deflac",
    "bilheteria_deflac",
    "outras_janelas_deflac",
    "receita_total_deflac",
    "roi_dom_total_deflac",
    "roi_internacional_0_100",
]:
    if col in obra:
        obra[col] = pd.to_numeric(obra[col], errors="coerce").fillna(0)

for col in [
    "investimento_fsa_deflac",
    "investimento_total_deflac",
    "bilheteria_deflac",
    "outras_janelas_deflac",
    "receita_total_deflac",
    "roi_dom_total_deflac",
    "roi_intl_medio",
    "roi_intl_max",
    "n_obras",
    "n_obras_fsa",
    "n_obras_renuncia",
]:
    if col in prod:
        prod[col] = pd.to_numeric(prod[col], errors="coerce").fillna(0)

for col in cham.columns:
    if col not in ["chamada", "categoria"]:
        cham[col] = pd.to_numeric(cham[col], errors="coerce").fillna(0)


def fmt_money(v: float) -> str:
    if pd.isna(v):
        return "n.d."
    v = float(v)
    if abs(v) >= 1e9:
        return f"R$ {v / 1e9:.2f} bi".replace(".", ",")
    if abs(v) >= 1e6:
        return f"R$ {v / 1e6:.1f} mi".replace(".", ",")
    if abs(v) >= 1e3:
        return f"R$ {v / 1e3:.0f} mil".replace(".", ",")
    return f"R$ {v:.0f}".replace(".", ",")


def fmt_num(v: float, digits: int = 0) -> str:
    if pd.isna(v):
        return "n.d."
    if digits == 0:
        return f"{int(round(float(v))):,}".replace(",", ".")
    return f"{float(v):,.{digits}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_pct(v: float, digits: int = 1) -> str:
    return f"{fmt_num(v, digits)}%"


def fmt_ratio(v: float) -> str:
    if pd.isna(v):
        return "n.d."
    return f"{float(v):.2f}x".replace(".", ",")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[list[str]], headers: list[str] | None = None) -> None:
    table = doc.add_table(rows=1 if headers else 0, cols=len(headers or rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if headers:
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_text(hdr[i], h, bold=True)
            set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def style_doc(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(68, 68, 68)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def savefig(name: str) -> Path:
    path = CHART_DIR / f"{name}.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close()
    return path


def fig_barh(df: pd.DataFrame, y: str, x: str, title: str, xlabel: str, name: str, color: str = COLORS["blue"]) -> Path:
    d = df.copy()
    plt.figure(figsize=(8.8, max(3.2, 0.34 * len(d))))
    plt.barh(d[y], d[x], color=color)
    plt.title(title, loc="left", fontsize=12, weight="bold")
    plt.xlabel(xlabel)
    plt.grid(axis="x", alpha=0.25)
    return savefig(name)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.65))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(8.5)


def classify_group(row: pd.Series) -> str:
    fsa = row["investimento_fsa_deflac"]
    ren = row["investimento_renuncia_total_deflac"]
    total = fsa + ren
    if fsa > 0 and ren <= 0:
        return "FSA puro"
    if ren > 0 and fsa <= 0:
        return "Renúncia pura"
    if fsa > 0 and ren > 0:
        return "Misto — FSA majoritário" if fsa / total >= 0.5 else "Misto — renúncia majoritária"
    return "Sem investimento público identificado"


obra["grupo_mecanismo"] = obra.apply(classify_group, axis=1)
obra["tem_roi_intl_13"] = obra["roi_internacional_0_100"] >= 13
prod["tem_roi_intl_13"] = prod["roi_intl_max"] >= 13
obra_ret = obra[obra["bilheteria_deflac"] > 0].copy()

prod_lookup = prod[["CNPJ_produtora", "razao_social", "UF", "cluster"]].copy()
prod_lookup["CNPJ_produtora"] = prod_lookup["CNPJ_produtora"].fillna("").astype(str)
prod_lookup = prod_lookup.drop_duplicates("CNPJ_produtora")
obra_ret["CNPJ_produtora"] = obra_ret["CNPJ_produtora"].fillna("").astype(str)
prod_comp = (
    obra_ret[obra_ret["CNPJ_produtora"].fillna("").astype(str).str.len() > 0]
    .groupby("CNPJ_produtora", as_index=False)
    .agg(
        n_obras=("CPB", "count"),
        investimento_fsa_deflac=("investimento_fsa_deflac", "sum"),
        investimento_total_deflac=("investimento_total_deflac", "sum"),
        bilheteria_deflac=("bilheteria_deflac", "sum"),
        outras_janelas_deflac=("outras_janelas_deflac", "sum"),
        receita_total_deflac=("receita_total_deflac", "sum"),
        roi_intl_max=("roi_internacional_0_100", "max"),
        roi_intl_medio=("roi_internacional_0_100", "mean"),
    )
    .merge(prod_lookup, on="CNPJ_produtora", how="left")
)
prod_comp["cluster"] = prod_comp["cluster"].fillna("Não classificado")
prod_comp["razao_social"] = prod_comp["razao_social"].fillna(prod_comp["CNPJ_produtora"])
prod_comp["UF"] = prod_comp["UF"].fillna("")
prod_comp["roi_dom_total_deflac"] = np.where(
    prod_comp["investimento_total_deflac"] > 0,
    prod_comp["receita_total_deflac"] / prod_comp["investimento_total_deflac"],
    0,
)
prod_comp["tem_roi_intl_13"] = prod_comp["roi_intl_max"] >= 13


def weighted_roi(g: pd.DataFrame) -> float:
    inv_sum = g["investimento_total_deflac"].sum()
    rec_sum = g["receita_total_deflac"].sum()
    return rec_sum / inv_sum if inv_sum > 0 else 0


group_summary = (
    obra_ret.groupby("grupo_mecanismo")
    .apply(
        lambda g: pd.Series(
            {
                "obras": len(g),
                "investimento": g["investimento_total_deflac"].sum(),
                "receita": g["receita_total_deflac"].sum(),
                "roi": weighted_roi(g),
                "pct_roi_intl_13": g["tem_roi_intl_13"].mean() * 100,
            }
        ),
        include_groups=False,
    )
    .reset_index()
)
group_summary = group_summary[group_summary["grupo_mecanismo"] != "Sem investimento público identificado"]

def load_panel_categories() -> pd.DataFrame:
    """Read the category metrics already produced by the category-selection panel."""
    html_path = ROOT / "resultados" / "painel_criterio_selecao.html"
    text = html_path.read_text(encoding="utf-8")
    match = re.search(r"const CATS\s*=\s*(\[.*?\]);", text, flags=re.S)
    if not match:
        raise RuntimeError(f"Nao foi possivel localizar const CATS em {html_path}")
    panel = pd.DataFrame(json.loads(match.group(1)))
    required = ["label", "n_obras", "inv", "bilh_prop", "roi_tot_def", "roi_fsa_def", "rda", "intl_avg", "intl_peak"]
    missing = [col for col in required if col not in panel.columns]
    if missing:
        raise RuntimeError(f"CATS do painel sem colunas esperadas: {missing}")
    return pd.DataFrame(
        {
            "categoria": panel["label"],
            "obras": pd.to_numeric(panel["n_obras"], errors="coerce").fillna(0),
            "investimento": pd.to_numeric(panel["inv"], errors="coerce").fillna(0),
            "receita": pd.to_numeric(panel["bilh_prop"], errors="coerce").fillna(0),
            "roi_agregado": pd.to_numeric(panel["roi_tot_def"], errors="coerce").fillna(0),
            "roi_fsa_def": pd.to_numeric(panel["roi_fsa_def"], errors="coerce").fillna(0),
            "roi_proporcional": pd.to_numeric(panel["rda"], errors="coerce").fillna(0),
            "roi_intl": pd.to_numeric(panel["intl_avg"], errors="coerce").fillna(0),
            "roi_intl_max": pd.to_numeric(panel["intl_peak"], errors="coerce").fillna(0),
            "roi_intl_total": pd.to_numeric(panel["intl_avg"], errors="coerce").fillna(0)
            * pd.to_numeric(panel["n_obras"], errors="coerce").fillna(0),
        }
    )


def load_concentration_panel_data() -> dict:
    html_path = ROOT / "resultados" / "painel_concentracao_produtoras.html"
    text = html_path.read_text(encoding="utf-8")
    match = re.search(r"const D\s*=\s*(\{.*?\});", text, flags=re.S)
    if not match:
        raise RuntimeError(f"Nao foi possivel localizar const D em {html_path}")
    return json.loads(match.group(1))


def load_panel_producer_clusters() -> pd.DataFrame:
    """Read the producer-cluster metrics from the same JSON used by the final panel."""
    html_path = ROOT / "resultados" / "painel_produtoras.html"
    text = html_path.read_text(encoding="utf-8")
    match = re.search(r"const PROD\s*=\s*(\[.*?\]);", text, flags=re.S)
    if not match:
        raise RuntimeError(f"Nao foi possivel localizar const PROD em {html_path}")

    panel = pd.DataFrame(json.loads(match.group(1)))
    required = ["cl", "n", "inv_def", "rec_def", "rim", "ria"]
    missing = [col for col in required if col not in panel.columns]
    if missing:
        raise RuntimeError(f"PROD do painel sem colunas esperadas: {missing}")

    cluster_names = {
        "duplo": "Duplo Retorno",
        "dom": "Retorno Doméstico",
        "intl": "Retorno Internacional",
        "sem_retorno": "Fomento Baixo Retorno",
        "pequeno": "Pequeno Porte",
    }
    panel = panel[panel["cl"].isin(cluster_names)].copy()
    for col in ["n", "inv_def", "rec_def", "rim", "ria"]:
        panel[col] = pd.to_numeric(panel[col], errors="coerce").fillna(0)
    panel["cluster"] = panel["cl"].map(cluster_names)
    panel["tem_roi_intl_13"] = panel["rim"] >= 13

    summary = (
        panel.groupby("cluster", as_index=False)
        .agg(
            produtoras=("cluster", "size"),
            obras=("n", "sum"),
            investimento=("inv_def", "sum"),
            receita=("rec_def", "sum"),
            roi_intl_max_med=("rim", "mean"),
            pct_roi_intl_13=("tem_roi_intl_13", "mean"),
        )
    )
    summary["roi_agregado"] = np.where(
        summary["investimento"] > 0,
        summary["receita"] / summary["investimento"],
        0,
    )
    summary["pct_roi_intl_13"] *= 100
    return summary


def load_panel_producers_for_plots() -> pd.DataFrame:
    """Normalize panel producer records to the column names used by report charts."""
    html_path = ROOT / "resultados" / "painel_produtoras.html"
    text = html_path.read_text(encoding="utf-8")
    match = re.search(r"const PROD\s*=\s*(\[.*?\]);", text, flags=re.S)
    if not match:
        raise RuntimeError(f"Nao foi possivel localizar const PROD em {html_path}")

    panel = pd.DataFrame(json.loads(match.group(1)))
    cluster_names = {
        "duplo": "Duplo Retorno",
        "dom": "Retorno Doméstico",
        "intl": "Retorno Internacional",
        "sem_retorno": "Fomento Baixo Retorno",
        "pequeno": "Pequeno Porte",
    }
    panel = panel[panel["cl"].isin(cluster_names)].copy()
    for col in ["n", "inv_def", "rec_def", "rim", "ria"]:
        panel[col] = pd.to_numeric(panel[col], errors="coerce").fillna(0)

    out = pd.DataFrame(
        {
            "cluster": panel["cl"].map(cluster_names),
            "n_obras": panel["n"],
            "investimento_total_deflac": panel["inv_def"],
            "receita_total_deflac": panel["rec_def"],
            "roi_intl_max": panel["rim"],
            "roi_intl_medio": panel["ria"],
        }
    )
    out["roi_dom_total_deflac"] = np.where(
        out["investimento_total_deflac"] > 0,
        out["receita_total_deflac"] / out["investimento_total_deflac"],
        0,
    )
    return out


def load_panel_curtas_uplift() -> dict[str, float] | None:
    """Read the short-to-feature uplift metrics from the final interactive panel."""
    panel_candidates = list((ROOT / "output_final").glob("An*lise do Retorno do Fomento P*blico ao Audiovisual Brasileiro*.html"))
    if not panel_candidates:
        return None
    text = panel_candidates[0].read_text(encoding="utf-8", errors="ignore")

    base = re.search(
        r"Na base geral,\s*<b[^>]*>(\d+)</b>\s*de\s*<b[^>]*>(\d+)</b>[\s\S]*?<b[^>]*>([\d.,]+)%</b>",
        text,
    )
    selected = re.search(
        r"Entre diretores com curta selecionado internacionalmente,\s*<b[^>]*>(\d+)</b>\s*de\s*<b[^>]*>(\d+)</b>[\s\S]*?<b[^>]*>([\d.,]+)%</b>",
        text,
    )
    gain = re.search(r"\+([\d.,]+)\s*p\.p\.\s*·\s*([\d.,]+)x", text)
    if not base or not selected:
        return None

    def pct(raw: str) -> float:
        raw = raw.strip()
        if "," in raw:
            return float(raw.replace(".", "").replace(",", "."))
        return float(raw)

    base_fest, base_total, base_pct = int(base.group(1)), int(base.group(2)), pct(base.group(3))
    sel_fest, sel_total, sel_pct = int(selected.group(1)), int(selected.group(2)), pct(selected.group(3))
    gain_pp = pct(gain.group(1)) if gain else sel_pct - base_pct
    gain_mult = pct(gain.group(2)) if gain else (sel_pct / base_pct if base_pct else 0)

    return {
        "base_fest": base_fest,
        "base_total": base_total,
        "base_pct": base_pct,
        "selected_fest": sel_fest,
        "selected_total": sel_total,
        "selected_pct": sel_pct,
        "gain_pp": gain_pp,
        "gain_mult": gain_mult,
    }


cat_fsa = load_panel_categories()
cat_summary = cat_fsa.copy()
concentration = load_concentration_panel_data()

cluster_summary = load_panel_producer_clusters()
prod_cluster = load_panel_producers_for_plots()
curtas_uplift = load_panel_curtas_uplift()


def gini(values: pd.Series) -> float:
    x = np.sort(pd.to_numeric(values, errors="coerce").fillna(0).to_numpy())
    x = x[x >= 0]
    if len(x) == 0 or x.sum() == 0:
        return 0
    n = len(x)
    return float((2 * np.arange(1, n + 1) @ x) / (n * x.sum()) - (n + 1) / n)


def concentration_rows(df: pd.DataFrame) -> list[list[str]]:
    d = df[df["investimento_fsa_deflac"] > 0].sort_values("investimento_fsa_deflac", ascending=False).copy()
    total = d["investimento_fsa_deflac"].sum()
    rows = []
    for label, frac in [("Top 1%", 0.01), ("Top 5%", 0.05), ("Top 10%", 0.10), ("Top 20%", 0.20), ("Metade inferior", 0.50)]:
        if label == "Metade inferior":
            sub = d.tail(max(1, math.ceil(len(d) * frac)))
        else:
            sub = d.head(max(1, math.ceil(len(d) * frac)))
        rows.append([label, fmt_num(len(sub)), fmt_pct(sub["investimento_fsa_deflac"].sum() / total * 100 if total else 0)])
    return rows


def make_charts() -> list[tuple[Path, str]]:
    charts: list[tuple[Path, str]] = []

    by_year = obra_ret.groupby("ano").agg(obras=("CPB", "count"), investimento=("investimento_total_deflac", "sum"), receita=("receita_total_deflac", "sum")).reset_index()
    fig, ax1 = plt.subplots(figsize=(8.8, 4.5))
    ax1.bar(by_year["ano"], by_year["investimento"] / 1e9, color=COLORS["blue"], alpha=0.75, label="Investimento")
    ax1.plot(by_year["ano"], by_year["receita"] / 1e9, color=COLORS["orange"], marker="o", label="Receita")
    ax1.set_title("Série anual: investimento público e receita estimada", loc="left", fontsize=12, weight="bold")
    ax1.set_ylabel("R$ bi, valores de 2024")
    ax1.grid(axis="y", alpha=0.25)
    ax1.legend(frameon=False)
    charts.append((savefig("01_serie_anual"), "Figura 1. Investimento total e receita estimada por ano de produção no recorte analítico."))

    d = group_summary.sort_values("investimento")
    charts.append((fig_barh(d, "grupo_mecanismo", "investimento", "Investimento por combinação de mecanismo", "R$ 2024", "02_investimento_grupo", COLORS["blue"]), "Figura 2. Distribuição do investimento entre FSA puro, renúncia pura e arranjos mistos."))

    d = group_summary.sort_values("roi")
    charts.append((fig_barh(d, "grupo_mecanismo", "roi", "ROI doméstico agregado por mecanismo", "Receita / investimento total", "03_roi_grupo", COLORS["green"]), "Figura 3. ROI doméstico agregado por grupo de mecanismo, com receita definida como bilheteria + outras janelas estimadas."))

    d = group_summary.sort_values("pct_roi_intl_13")
    charts.append((fig_barh(d, "grupo_mecanismo", "pct_roi_intl_13", "Obras com ROI internacional qualificado", "% de obras com score >= 13", "04_pct_intl_grupo", COLORS["purple"]), "Figura 4. Percentual de obras com ROI Internacional >= 13 por grupo de mecanismo."))

    d = cat_fsa.sort_values("investimento").tail(12)
    charts.append((fig_barh(d, "categoria", "investimento", "Categorias FSA: investimento total", "R$ 2024", "05_categoria_investimento", COLORS["blue"]), "Figura 5. Categorias FSA com maior volume de investimento total deflacionado."))

    d = cat_fsa[cat_fsa["obras"] >= 10].sort_values("roi_intl_total").tail(12).copy()
    y = np.arange(len(d))
    fig, ax1 = plt.subplots(figsize=(9.2, 5.2))
    ax1.barh(y, d["roi_intl_total"], color=COLORS["purple"], alpha=0.78, label="Score total")
    ax1.set_yticks(y, d["categoria"])
    ax1.set_xlabel("Score internacional total (media x obras)")
    ax1.grid(axis="x", alpha=0.25)
    ax2 = ax1.twiny()
    ax2.plot(d["roi_intl"], y, color=COLORS["orange"], marker="o", linewidth=2, label="Score medio")
    ax2.set_xlabel("Score internacional medio (0-100)")
    ax1.set_title("Categorias FSA: score internacional total e medio", loc="left", fontsize=12, weight="bold")
    lines, labels = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines + lines2, labels + labels2, loc="lower right", frameon=False)
    charts.append((savefig("07_categoria_intl"), "Figura 7. Score internacional total e médio por categoria FSA. O total corresponde ao score médio multiplicado pelo número de obras da categoria."))

    d = cat_fsa[cat_fsa["obras"] >= 10].sort_values("roi_agregado").tail(12)
    charts.append((fig_barh(d, "categoria", "roi_agregado", "Categorias FSA: ROI total deflacionado", "Receita proporcional / capital publico total", "06_categoria_roi", COLORS["green"]), "Figura 6. ROI Total Deflacionado por categoria FSA, conforme a metodologia do painel, restringindo categorias com pelo menos dez obras."))

    d = cluster_summary.sort_values("produtoras")
    plt.figure(figsize=(9.2, 4.9))
    bars = plt.barh(d["cluster"], d["produtoras"], color=[CLUSTER_COLORS.get(c, COLORS["gray"]) for c in d["cluster"]])
    plt.xlabel("N produtoras")
    plt.title("Produtoras por cluster", loc="left", fontsize=12, weight="bold")
    plt.grid(axis="x", alpha=0.25)
    for bar, value in zip(bars, d["produtoras"]):
        plt.text(bar.get_width() + max(d["produtoras"]) * 0.01, bar.get_y() + bar.get_height() / 2, fmt_num(value), va="center", fontsize=9)
    charts.append((savefig("08_cluster_counts"), "Figura 8. Distribuição de produtoras independentes por cluster, com retorno internacional definido por ROI intl máximo >= 13."))

    order = ["Duplo Retorno", "Retorno Doméstico", "Retorno Internacional", "Fomento Baixo Retorno", "Pequeno Porte"]
    d = cluster_summary.set_index("cluster").reindex([x for x in order if x in cluster_summary["cluster"].values]).reset_index()
    y = np.arange(len(d))
    height = 0.36
    plt.figure(figsize=(9.2, 4.9))
    plt.barh(y - height / 2, d["investimento"] / 1e9, height, label="Investimento", color=COLORS["blue"])
    plt.barh(y + height / 2, d["receita"] / 1e9, height, label="Receita", color=COLORS["orange"])
    plt.yticks(y, d["cluster"])
    plt.xlabel("R$ bi, valores de 2024")
    plt.title("Investimento e receita por cluster de produtora", loc="left", fontsize=12, weight="bold")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(frameon=False)
    charts.append((savefig("09_cluster_inv_receita"), "Figura 9. Investimento e receita estimada por cluster de produtora."))

    plt.figure(figsize=(9.2, 4.9))
    height = 0.36
    plt.barh(y - height / 2, d["roi_agregado"], height, color=[CLUSTER_COLORS.get(c, COLORS["gray"]) for c in d["cluster"]], label="ROI doméstico")
    plt.barh(y + height / 2, d["roi_intl_max_med"] / 13, height, color=COLORS["purple"], alpha=0.72, label="ROI intl médio / 13")
    plt.axvline(1, color="#777777", linestyle="--", linewidth=1, alpha=0.7)
    plt.yticks(y, d["cluster"])
    plt.xlabel("ROI doméstico (x) e ROI intl médio dividido pelo corte 13")
    plt.title("Retorno doméstico e intensidade internacional por cluster", loc="left", fontsize=12, weight="bold")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(frameon=False)
    charts.append((savefig("10_cluster_roi"), "Figura 10. ROI doméstico agregado e intensidade internacional média por cluster; a linha pontilhada marca 1x ou o corte internacional normalizado."))

    dplot = prod_cluster[(prod_cluster["investimento_total_deflac"] > 0) & (prod_cluster["receita_total_deflac"] > 0)].copy()
    plt.figure(figsize=(8.8, 5.2))
    for cl, g in dplot.groupby("cluster"):
        plt.scatter(g["investimento_total_deflac"], g["receita_total_deflac"], s=np.clip(g["n_obras"] * 8, 12, 180), alpha=0.55, label=cl, color=CLUSTER_COLORS.get(cl, COLORS["gray"]))
    plt.xscale("log")
    plt.yscale("log")
    plt.xlabel("Investimento total deflacionado, R$ 2024 (log)")
    plt.ylabel("Receita total estimada, R$ 2024 (log)")
    plt.title("Produtoras: investimento x receita", loc="left", fontsize=12, weight="bold")
    plt.grid(alpha=0.25)
    plt.legend(frameon=False, fontsize=7)
    charts.append((savefig("11_scatter_produtoras"), "Figura 11. Relação entre investimento e receita por produtora; tamanho do ponto representa número de obras."))

    d = prod_comp[prod_comp["investimento_fsa_deflac"] > 0]["investimento_fsa_deflac"].sort_values().to_numpy()
    cum = np.cumsum(d) / d.sum() if d.sum() else np.zeros_like(d)
    x = np.arange(1, len(d) + 1) / len(d)
    plt.figure(figsize=(6.8, 5.2))
    plt.plot(x, cum, color=COLORS["blue"], linewidth=2.5, label="FSA observado")
    plt.plot([0, 1], [0, 1], "--", color="#999999", label="igualdade perfeita")
    plt.title("Curva de Lorenz do investimento FSA por produtora", loc="left", fontsize=12, weight="bold")
    plt.xlabel("Participação acumulada das produtoras")
    plt.ylabel("Participação acumulada do FSA")
    plt.grid(alpha=0.25)
    plt.legend(frameon=False)
    charts.append((savefig("12_lorenz"), f"Figura 12. Concentração do FSA entre produtoras com obras de bilheteria positiva e investimento FSA positivo; Gini = {gini(prod_comp['investimento_fsa_deflac']):.3f}."))

    tier = pd.DataFrame(concentration["tier_rows"])
    tier = tier.iloc[::-1].copy()
    y = np.arange(len(tier))
    fig, ax1 = plt.subplots(figsize=(9.2, 5.0))
    ax1.barh(y, tier["fsa_share"], color=tier["color"], alpha=0.82, label="% do FSA")
    ax1.set_yticks(y, tier["name"])
    ax1.set_xlabel("% do FSA nominal")
    ax1.grid(axis="x", alpha=0.25)
    ax2 = ax1.twiny()
    ax2.plot(tier["ticket_med"], y, color=COLORS["orange"], marker="o", linewidth=2.2, label="Ticket mediano anual")
    ax2.set_xlabel("Ticket mediano anual, R$ mil")
    ax1.set_title("Concentração por tier de produtora", loc="left", fontsize=12, weight="bold")
    lines, labels = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines + lines2, labels + labels2, frameon=False, loc="lower right")
    charts.append((savefig("20_concentracao_tiers"), "Figura 13. Participação no FSA e ticket mediano anual por tier de produtora, conforme painel de concentração."))

    coverage = pd.DataFrame(
        {
            "indicador": ["Bilheteria", "Outras janelas", "ROI intl >=13", "Festival", "Lumière", "VOD Europa", "Crítica >=2 fontes"],
            "obras": [
                (obra["bilheteria_deflac"] > 0).sum(),
                (obra["outras_janelas_deflac"] > 0).sum(),
                (obra["roi_internacional_0_100"] >= 13).sum(),
                (obra["pontuacao_festivais"] > 0).sum(),
                (obra["adm_eu_lumiere"] > 0).sum(),
                (obra["vod_n_paises"] > 0).sum(),
                (obra["critica_n_fontes"] >= 2).sum(),
            ],
        }
    ).sort_values("obras")
    charts.append((fig_barh(coverage, "indicador", "obras", "Cobertura empírica por tipo de evidência", "N obras", "13_cobertura", COLORS["teal"]), "Figura A1. Cobertura de dados por fonte de evidência na base de obras."))

    crit = obra[(obra["critica_indice_1_5"] > 0) & (obra["critica_n_fontes"] >= 2)].copy()
    if len(crit) > 5:
        plt.figure(figsize=(8.8, 4.8))
        plt.scatter(crit["critica_indice_1_5"], crit["roi_internacional_0_100"], alpha=0.45, s=24, color=COLORS["purple"])
        plt.xlabel("Índice crítico, 1-5")
        plt.ylabel("ROI Internacional, 0-100")
        plt.title("Crítica especializada e circulação internacional", loc="left", fontsize=12, weight="bold")
        plt.grid(alpha=0.25)
        charts.append((savefig("14_critica_intl"), "Figura 14. Relação exploratória entre índice crítico e score internacional; não representa inferência causal."))

    uf = prod_comp.groupby("UF").agg(produtoras=("UF", "size"), investimento=("investimento_total_deflac", "sum")).reset_index()
    uf = uf[uf["UF"].astype(str).str.len() > 0].sort_values("investimento").tail(12)
    charts.append((fig_barh(uf, "UF", "investimento", "Investimento por UF da produtora", "R$ 2024", "15_uf_investimento", COLORS["blue"]), "Figura 15. Distribuição territorial do investimento total por UF da produtora."))

    raca = pd.DataFrame({"etapa": ["Inscritos sem PA", "Inscritos com PA", "Selecionados com PA"], "pct": [15.2, 22.8, 32.4]})
    charts.append((fig_barh(raca.sort_values("pct"), "etapa", "pct", "Raça: participação de pessoas negras", "%", "16_genero_direcao", COLORS["teal"]), "Figura 16. Inscritos e selecionados por raça: percentual de pessoas negras nas etapas reportadas pelo painel."))

    genero = pd.DataFrame({"etapa": ["Inscritas com PA", "Selecionadas com PA"], "pct": [37.0, 52.6]})
    charts.append((fig_barh(genero.sort_values("pct"), "etapa", "pct", "Gênero: participação de mulheres", "%", "17_diversidade_regional", COLORS["orange"]), "Figura 17. Inscritos e selecionados por gênero: percentual de mulheres nas etapas reportadas pelo painel."))

    cf = curtas.groupby("festival").agg(selecoes=("titulo", "count"), premiados=("premiado", "sum")).reset_index().sort_values("selecoes")
    plt.figure(figsize=(8.8, max(3.2, 0.42 * len(cf))))
    plt.barh(cf["festival"], cf["selecoes"], color=COLORS["purple"], label="seleções")
    plt.barh(cf["festival"], cf["premiados"], color=COLORS["yellow"], label="premiados")
    plt.title("Curtas brasileiros em festivais internacionais", loc="left", fontsize=12, weight="bold")
    plt.xlabel("N registros")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(frameon=False)
    charts.append((savefig("18_curtas_festivais"), "Figura 18. Seleções e premiações de curtas brasileiros em festivais internacionais mapeados."))

    cit = citacoes.copy()
    cit["CITA_SOMA_CIT"] = pd.to_numeric(cit["CITA_SOMA_CIT"], errors="coerce").fillna(0)
    top_cit = cit[cit["CITA_SOMA_CIT"] > 0].nlargest(15, "CITA_SOMA_CIT").sort_values("CITA_SOMA_CIT")
    charts.append((fig_barh(top_cit, "DIRETOR", "CITA_SOMA_CIT", "Diretores mais citados em bases acadêmicas", "Citações OpenAlex", "19_citacoes_diretores", COLORS["green"]), "Figura 19. Diretores com maior soma de citações acadêmicas no enriquecimento OpenAlex."))

    # ── NOVO: Vocação Comercial vs. Alcance Internacional (quadrante por categoria) ────
    d_quad = cat_fsa[cat_fsa["obras"] >= 5].copy()
    if len(d_quad) >= 3:
        d_quad["bubble"] = np.clip(d_quad["investimento"] / 8e6, 40, 900)
        fig, ax = plt.subplots(figsize=(10.5, 6.5))
        ax.scatter(d_quad["roi_agregado"], d_quad["roi_intl"], s=d_quad["bubble"],
                   alpha=0.70, color=COLORS["blue"], edgecolors="white", linewidths=0.6)
        for _, row in d_quad.iterrows():
            label = (row["categoria"]
                     .replace("FSA Pontuação ", "Pont.")
                     .replace("FSA Automático ", "Auto.")
                     .replace("FSA Complementação", "Complement.")
                     .replace("FSA Comercialização / Distribuição", "Comerc./Dist.")
                     .replace("FSA Coprodução Internacional", "Coprod. Intl.")
                     .replace(" e Roteiro", ""))
            ax.annotate(label, (row["roi_agregado"], row["roi_intl"]),
                        fontsize=7.5, ha="center", va="bottom",
                        xytext=(0, 7), textcoords="offset points")
        med_x = d_quad["roi_agregado"].median()
        med_y = d_quad["roi_intl"].median()
        ax.axvline(x=med_x, color="#aaaaaa", linestyle="--", linewidth=1.1, alpha=0.75)
        ax.axhline(y=med_y, color="#aaaaaa", linestyle="--", linewidth=1.1, alpha=0.75)
        ax.set_xlabel("ROI Total Deflacionado (receita proporcional / capital público total)", fontsize=10)
        ax.set_ylabel("ROI Internacional Médio (score composto 0–100)", fontsize=10)
        ax.set_title("Vocação Comercial vs. Alcance Internacional por Categoria FSA",
                     loc="left", fontsize=12, weight="bold")
        ax.grid(alpha=0.18)
        plt.tight_layout()
        charts.append((savefig("21_quadrant_categorias"),
                       "Figura 21. Posicionamento de cada categoria FSA no espaço retorno doméstico × alcance "
                       "internacional. Tamanho proporcional ao investimento total. Linhas: medianas do conjunto."))

    # ── NOVO: Scatter ROI Dom × ROI Intl por produtora (matriz de portfolio) ───────────
    dsc = prod_cluster[prod_cluster["investimento_total_deflac"] > 0].copy()
    if "roi_dom_total_deflac" in dsc.columns:
        roi_dom_col = "roi_dom_total_deflac"
    else:
        dsc["_roi_dom"] = np.where(dsc["investimento_total_deflac"] > 0,
                                   dsc["receita_total_deflac"] / dsc["investimento_total_deflac"], 0)
        roi_dom_col = "_roi_dom"
    dsc_plot = dsc[dsc[roi_dom_col] < 6].copy()
    fig, ax = plt.subplots(figsize=(10, 6))
    for cl, g in dsc_plot.groupby("cluster"):
        ax.scatter(g[roi_dom_col], g["roi_intl_medio"],
                   s=np.clip(g["n_obras"] * 7, 10, 200),
                   alpha=0.52, label=cl, color=CLUSTER_COLORS.get(cl, COLORS["gray"]),
                   edgecolors="white", linewidths=0.4)
    ax.axvline(x=1.0, color="#888888", linestyle="--", linewidth=1.1, alpha=0.65)
    ax.axhline(y=13, color="#888888", linestyle="--", linewidth=1.1, alpha=0.65)
    ax.set_xlabel("ROI Doméstico Total Deflacionado", fontsize=10)
    ax.set_ylabel("ROI Internacional Médio (score 0–100)", fontsize=10)
    ax.set_title("Matriz de Portfolio — ROI Doméstico vs. ROI Internacional por Produtora",
                 loc="left", fontsize=12, weight="bold")
    ax.legend(frameon=False, fontsize=8, loc="upper right")
    ax.grid(alpha=0.18)
    plt.tight_layout()
    charts.append((savefig("22_scatter_roi_produtoras"),
                   "Figura 22. Cada ponto é uma produtora; cor = cluster. "
                   "Linhas tracejadas: ROI Dom = 1,0x e ROI Intl = 13 (limiares analíticos)."))

    # ── NOVO: Distribuição acumulada do FSA por produtora (curva de threshold) ──────────
    d_tick = (prod_comp[prod_comp["investimento_fsa_deflac"] > 0]["investimento_fsa_deflac"]
              .sort_values().values)
    if len(d_tick) > 10:
        cum_pct = np.arange(1, len(d_tick) + 1) / len(d_tick) * 100
        fig, ax = plt.subplots(figsize=(8.8, 5.0))
        ax.plot(d_tick / 1e6, cum_pct, color=COLORS["blue"], linewidth=2.3)
        for thresh, color, label in [
            (0.5, COLORS["red"],    "R$ 500 mil"),
            (1.0, COLORS["orange"], "R$ 1 mi"),
            (2.0, COLORS["green"],  "R$ 2 mi"),
        ]:
            pct_at = (d_tick <= thresh * 1e6).mean() * 100
            ax.axvline(x=thresh, color=color, linestyle="--", linewidth=1.4,
                       alpha=0.85, label=f"{label} ({pct_at:.0f}% abaixo)")
        ax.set_xlabel("FSA total recebido pela produtora, R$ mi (escala log)")
        ax.set_ylabel("% acumulada das produtoras com FSA positivo")
        ax.set_xscale("log")
        ax.set_title("Distribuição acumulada do FSA por produtora",
                     loc="left", fontsize=12, weight="bold")
        ax.legend(frameon=False)
        ax.grid(alpha=0.22)
        plt.tight_layout()
        charts.append((savefig("23_ticket_acumulado"),
                       "Figura 23. Curva acumulada do FSA total por produtora com investimento positivo. "
                       "Linhas verticais marcam limiares de R$ 500 mil, R$ 1 mi e R$ 2 mi com % abaixo."))

    # ── NOVO: Taxas de seleção com/sem PA — diversidade (raça + gênero) ─────────────────
    pa_raca = pd.DataFrame({
        "grupo":  ["Negros\nsem PA", "Brancos\nsem PA", "Negros\ncom PA", "Brancos\ncom PA"],
        "taxa":   [26.8, 29.9, 14.8, 9.1],
        "cor":    [COLORS["orange"], COLORS["gray"], COLORS["green"], COLORS["blue"]],
    })
    pa_gen = pd.DataFrame({
        "grupo":  ["Mulheres\ncom PA", "Homens\ncom PA"],
        "taxa":   [15.0, 7.9],
        "cor":    [COLORS["purple"], COLORS["blue"]],
    })
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11.5, 4.8))
    bars1 = ax1.bar(pa_raca["grupo"], pa_raca["taxa"], color=pa_raca["cor"], alpha=0.82, width=0.55)
    ax1.set_title("Raça: taxa de seleção FSA (%)", loc="left", fontsize=11, weight="bold")
    ax1.set_ylabel("% selecionados / inscritos")
    ax1.set_ylim(0, max(pa_raca["taxa"]) * 1.18)
    ax1.grid(axis="y", alpha=0.22)
    for bar in bars1:
        ax1.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                 f"{bar.get_height():.1f}%", ha="center", va="bottom", fontsize=9.5)
    bars2 = ax2.bar(pa_gen["grupo"], pa_gen["taxa"], color=pa_gen["cor"], alpha=0.82, width=0.35)
    ax2.set_title("Gênero: taxa de seleção FSA, editais com PA (%)", loc="left", fontsize=11, weight="bold")
    ax2.set_ylabel("% selecionados / inscritos")
    ax2.set_ylim(0, max(pa_gen["taxa"]) * 1.18)
    ax2.grid(axis="y", alpha=0.22)
    for bar in bars2:
        ax2.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
                 f"{bar.get_height():.1f}%", ha="center", va="bottom", fontsize=9.5)
    plt.tight_layout()
    charts.append((savefig("24_diversidade_taxas"),
                   "Figura 24. Taxas de seleção por raça (com e sem política afirmativa) e por gênero "
                   "(editais com PA). Dados: BRDE/FSA 2015–2025, 99 editais, função Diretor."))

    # ── NOVO: Índice crítico médio por categoria FSA ────────────────────────────────────
    if "critica_media" in cham.columns:
        crit_cat = (cham.groupby("categoria")
                    .agg(critica=("critica_media", "mean"), obras=("n_obras", "sum"))
                    .reset_index())
        crit_cat = crit_cat[crit_cat["obras"] >= 5].sort_values("critica")
        if len(crit_cat) >= 3:
            plt.figure(figsize=(9.5, max(3.5, 0.40 * len(crit_cat))))
            plt.barh(crit_cat["categoria"], crit_cat["critica"],
                     color=COLORS["yellow"], alpha=0.85)
            media_geral = crit_cat["critica"].mean()
            plt.axvline(x=media_geral, color=COLORS["red"], linestyle="--",
                        linewidth=1.4, alpha=0.85,
                        label=f"Média geral ({media_geral:.2f})")
            plt.xlabel("Índice crítico médio (escala 1–5)")
            plt.title("Qualidade crítica média por categoria FSA",
                      loc="left", fontsize=12, weight="bold")
            plt.grid(axis="x", alpha=0.22)
            plt.legend(frameon=False)
            plt.tight_layout()
            charts.append((savefig("25_critica_categoria"),
                           "Figura 25. Índice crítico médio (1–5, múltiplas fontes) por categoria FSA. "
                           "Linha tracejada = média geral das categorias com >= 5 obras."))

    # ── NOVO: Volume anual de seleções de curtas (se coluna de ano disponível) ──────────
    curtas_ano_col = next((c for c in ["ano", "Ano", "year", "Year", "data", "Data"]
                           if c in curtas.columns), None)
    if curtas_ano_col:
        try:
            ca = curtas.copy()
            ca["_ano"] = pd.to_numeric(ca[curtas_ano_col], errors="coerce")
            ca = ca.dropna(subset=["_ano"])
            ca["_ano"] = ca["_ano"].astype(int)
            by_ano_c = (ca.groupby("_ano")
                        .agg(selecoes=("titulo", "count"),
                             premiados=("premiado", "sum"))
                        .reset_index())
            if len(by_ano_c) >= 3:
                fig, ax = plt.subplots(figsize=(9.5, 4.2))
                ax.bar(by_ano_c["_ano"], by_ano_c["selecoes"],
                       color=COLORS["purple"], alpha=0.72, label="Seleções")
                ax.bar(by_ano_c["_ano"], by_ano_c["premiados"],
                       color=COLORS["yellow"], alpha=0.88, label="Premiações")
                ax.set_title("Curtas brasileiros em festivais internacionais — volume anual",
                             loc="left", fontsize=12, weight="bold")
                ax.set_xlabel("Ano")
                ax.set_ylabel("N registros")
                ax.legend(frameon=False)
                ax.grid(axis="y", alpha=0.22)
                plt.tight_layout()
                charts.append((savefig("26_curtas_por_ano"),
                               "Figura 26. Volume anual de seleções e premiações de curtas brasileiros "
                               "nos festivais internacionais mapeados."))
        except Exception:
            pass

    if curtas_uplift:
        labels = ["Base geral de\ndiretores de longas", "Diretores com\ncurta selecionado"]
        values = [curtas_uplift["base_pct"], curtas_uplift["selected_pct"]]
        counts = [
            f"{int(curtas_uplift['base_fest'])}/{int(curtas_uplift['base_total'])}",
            f"{int(curtas_uplift['selected_fest'])}/{int(curtas_uplift['selected_total'])}",
        ]
        fig, ax = plt.subplots(figsize=(8.8, 4.8))
        bars = ax.bar(labels, values, color=[COLORS["gray"], COLORS["teal"]], width=0.58)
        ax.set_title("Pergunta comparativa: quanto o curta selecionado aumenta a chance?",
                     loc="left", fontsize=12, weight="bold")
        ax.set_ylabel("% com longa em festival internacional")
        ax.set_ylim(0, max(values) * 1.42 + 4)
        ax.grid(axis="y", alpha=0.25)
        for bar, value, count in zip(bars, values, counts):
            ax.text(bar.get_x() + bar.get_width() / 2, value + max(values) * 0.04,
                    f"{count}\n{fmt_pct(value)}", ha="center", va="bottom",
                    fontsize=10, weight="bold")
        ax.annotate(
            f"+{fmt_num(curtas_uplift['gain_pp'], 1)} p.p. | {fmt_num(curtas_uplift['gain_mult'], 1)}x",
            xy=(1, values[1]),
            xytext=(0.56, values[1] + max(values) * 0.23),
            arrowprops=dict(arrowstyle="->", color=COLORS["orange"], linewidth=1.4),
            color=COLORS["orange"],
            fontsize=11,
            weight="bold",
        )
        charts.append((savefig("27_curtas_uplift"),
                       "Figura 27. Comparação entre a taxa de chegada a festivais internacionais de longas "
                       "na base geral de diretores e no recorte de diretores com curta selecionado."))

    return charts


charts = make_charts()
chart = {path.stem: (path, caption) for path, caption in charts}

# ── Pré-cálculos para parágrafos analíticos ───────────────────────────────────
by_year = (
    obra_ret.groupby("ano")
    .agg(obras=("CPB", "count"),
         investimento=("investimento_total_deflac", "sum"),
         receita=("receita_total_deflac", "sum"))
    .reset_index()
)
peak_inv_year = int(by_year.loc[by_year["investimento"].idxmax(), "ano"])
peak_inv_val  = by_year["investimento"].max()
peak_rec_year = int(by_year.loc[by_year["receita"].idxmax(), "ano"])
n_anos_serie  = by_year["ano"].nunique()

top_grupo      = group_summary.sort_values("investimento", ascending=False).iloc[0]
best_roi_grupo = group_summary.sort_values("roi", ascending=False).iloc[0]
worst_roi_grupo= group_summary.sort_values("roi").iloc[0]

top_cat_inv  = cat_fsa.sort_values("investimento", ascending=False).iloc[0] if len(cat_fsa) else None
cat_10       = cat_fsa[cat_fsa["obras"] >= 10]
best_cat_roi = cat_10.sort_values("roi_agregado", ascending=False).iloc[0] if len(cat_10) else None
best_cat_intl= cat_10.sort_values("roi_intl", ascending=False).iloc[0] if len(cat_10) else None

top_cluster    = cluster_summary.sort_values("roi_agregado", ascending=False).iloc[0] if len(cluster_summary) else None
intl_cluster   = cluster_summary.sort_values("roi_intl_max_med", ascending=False).iloc[0] if len(cluster_summary) else None
biggest_cluster= cluster_summary.sort_values("produtoras", ascending=False).iloc[0] if len(cluster_summary) else None

total_inv  = obra_ret["investimento_total_deflac"].sum()
total_rec  = obra_ret["receita_total_deflac"].sum()
roi_global = total_rec / total_inv if total_inv else 0
intl_works = int((obra_ret["roi_internacional_0_100"] >= 13).sum())
prod_intl  = int((prod_comp["roi_intl_max"] >= 13).sum())

gini_all  = gini(prod_comp["investimento_fsa_deflac"])
gini_pos  = gini(prod_comp.loc[prod_comp["investimento_fsa_deflac"] > 0, "investimento_fsa_deflac"])
n_pos_fsa = int((prod_comp["investimento_fsa_deflac"] > 0).sum())
n_below_500k = int(prod_comp["investimento_fsa_deflac"].between(1, 500_000).sum())
top10_fsa_share = (prod_comp.sort_values("investimento_fsa_deflac", ascending=False)
                   .head(10)["investimento_fsa_deflac"].sum()
                   / prod_comp["investimento_fsa_deflac"].sum() * 100
                   if prod_comp["investimento_fsa_deflac"].sum() > 0 else 0)

uf_inv = (prod_comp.groupby("UF")
          .agg(investimento=("investimento_total_deflac", "sum"))
          .reset_index())
uf_inv = uf_inv[uf_inv["UF"].astype(str).str.len() > 0].sort_values("investimento", ascending=False)
top_uf = uf_inv.iloc[0] if len(uf_inv) > 0 else None
top2_uf_share = uf_inv.head(2)["investimento"].sum() / uf_inv["investimento"].sum() * 100 if uf_inv["investimento"].sum() > 0 else 0

crit_obras = obra[(obra["critica_indice_1_5"] > 0) & (obra["critica_n_fontes"] >= 2)].copy()

# ── Documento ─────────────────────────────────────────────────────────────────
doc = Document()
style_doc(doc)

title = doc.add_paragraph()
title.style = doc.styles["Title"]
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Análise de Dados do Fomento Público ao Audiovisual Brasileiro").bold = True
subtitle = doc.add_paragraph("Relatório técnico enriquecido com gráficos, metodologia de cruzamento e ressalvas analíticas")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)
doc.add_paragraph()
meta = doc.add_paragraph(
    "Base consolidada do projeto · valores monetários deflacionados para R$ 2024 · "
    "classificação internacional de produtoras por ROI Internacional máximo >= 13"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()


doc.add_heading("Sumário executivo", level=1)
total_inv = obra_ret["investimento_total_deflac"].sum()
total_rec = obra_ret["receita_total_deflac"].sum()
roi_global = total_rec / total_inv if total_inv else 0
intl_works = int((obra_ret["roi_internacional_0_100"] >= 13).sum())
prod_intl = int((prod_comp["roi_intl_max"] >= 13).sum())
gini_all = gini(prod_comp["investimento_fsa_deflac"])
gini_pos = gini(prod_comp.loc[prod_comp["investimento_fsa_deflac"] > 0, "investimento_fsa_deflac"])

summary_bullets = [
    f"A base consolidada contém {fmt_num(len(obra))} obras entre {int(obra['ano'].min())} e {int(obra['ano'].max())}; o recorte de comparação de retorno usa {fmt_num(len(obra_ret))} obras com bilheteria positiva.",
    f"Dentro desse recorte comparável, o investimento público total deflacionado soma {fmt_money(total_inv)} e a receita estimada soma {fmt_money(total_rec)}, resultando em ROI doméstico agregado de {fmt_ratio(roi_global)}.",
    f"A classificação de retorno internacional para empresas foi endurecida: uma produtora só é tratada como internacional quando seu ROI Internacional máximo é >= 13. No recorte com bilheteria positiva, {fmt_num(prod_intl)} produtoras têm retorno internacional qualificado.",
    f"A concentração permanece material no recorte comparável: Gini do FSA entre produtoras com obras de bilheteria positiva = {fmt_num(gini_all, 3)}; entre produtoras com FSA positivo = {fmt_num(gini_pos, 3)}.",
    "As estimativas descrevem associações observadas na base consolidada. Elas não identificam efeito causal dos mecanismos de fomento sem desenho contrafactual adicional.",
]
for b in summary_bullets:
    doc.add_paragraph(b, style="List Bullet")

add_table(
    doc,
    [
        ["Obras na base consolidada", fmt_num(len(obra))],
        ["Obras no recorte comparável (bilheteria > 0)", fmt_num(len(obra_ret))],
        ["Produtoras no recorte comparável", fmt_num(len(prod_comp))],
        ["Chamadas/categorias FSA agregadas", fmt_num(len(cham))],
        ["Investimento total deflacionado", fmt_money(total_inv)],
        ["Receita total estimada", fmt_money(total_rec)],
        ["ROI doméstico agregado", fmt_ratio(roi_global)],
        ["Obras com ROI Internacional >= 13 no recorte", fmt_num(intl_works)],
        ["Produtoras com ROI Internacional máx >= 13 no recorte", fmt_num(prod_intl)],
    ],
    headers=["Indicador", "Valor"],
)

doc.add_heading("1. Escopo, unidades de análise e arquivos usados", level=1)
doc.add_paragraph(
    "O relatório foi reconstruído a partir dos artefatos consolidados da pipeline. A unidade de análise muda conforme a pergunta: "
    "obras para desempenho e cobertura, investimentos para decomposição por mecanismo, chamadas/categorias para comparação entre linhas FSA, "
    "e produtoras para concentração, capacidade empresarial e clusters."
)
add_table(
    doc,
    [
        ["base_nivel_obra.csv", fmt_num(len(obra)), "Obra/CPB", "Retorno doméstico, score internacional, crítica, direção e atributos da obra."],
        ["base_nivel_investimento.csv", fmt_num(len(inv)), "Obra x tipo de investimento", "Decomposição FSA/renúncia e ROI sobre cada parcela de investimento."],
        ["base_nivel_chamada.csv", fmt_num(len(cham)), "Chamada/categoria", "Agregados de investimento, receita, ROI e alcance por categoria FSA."],
        ["base_nivel_produtora.csv", fmt_num(len(prod)), "Produtora/CNPJ", "Clusters, concentração, retorno doméstico e internacional por empresa."],
        ["tabela_consolidada_obras.xlsx", fmt_num(len(master)), "Obra consolidada", "Base-mãe da pipeline antes dos recortes analíticos."],
    ],
    headers=["Arquivo", "Linhas", "Unidade", "Uso principal"],
)

doc.add_heading("2. Descrição dos dados", level=1)
doc.add_paragraph(
    "A base combina fontes administrativas, bases de mercado e evidências de circulação internacional. "
    "Os valores monetários usados na análise principal estão em reais de 2024. A receita total estimada é a soma de bilheteria deflacionada e estimativas de outras janelas."
)
add_table(
    doc,
    [
        ["FSA direto", "raw/projetos-fsa.csv; raw/obras-nao-pub-brasileiras-investimento-fsa.csv", "Valores aprovados/identificados por projeto, chamada e obra."],
        ["Renúncia fiscal", "raw/projetos-com-renuncia-fiscal.csv; raw/obras-nao-pub-brasileiras-fomento-indireto.csv", "Valores de Art. 3/3-A/39 e demais mecanismos indiretos."],
        ["Obras e metadados", "raw/obras-nao-pub-brasileiras-csv/*.csv", "CPB, título, ano, atributos de obra e dados de catalogação."],
        ["Produtoras e agentes", "raw/produtores-de-obras-nao-publicitarias-brasileiras.csv; raw/agentes-economicos-regulares.csv", "CNPJ, razão social, UF, classificação do agente e vínculo obra-produtora."],
        ["Bilheteria", "raw/bilheteria_brasileira_consolidada.xlsx; raw/bilheteria-agregado/por_filme_ano.csv", "Público e renda de salas de cinema, consolidados por obra."],
        ["Festivais", "resultados/festivais_consolidado.csv; dados/festivais_por_obra_ata_fsa2024.csv; tabelas_apoio/Festivais_por_obra_pre_expansao.xlsx", "Pontuação de circulação e prestígio em festivais."],
        ["Lumière/VOD Europa", "raw/lumiere_search.xlsx; raw/lumiere_vod_search.xlsx", "Admissões europeias e presença em plataformas/países VOD."],
        ["Crítica e direção", "dados/critica_obras.csv; dados/citacoes_diretores.csv; dados/prestigio_diretores.csv; dados/imdb_enrichment.csv", "Índice crítico, citações acadêmicas, prestígio de diretores e enriquecimentos auxiliares."],
        ["Deflação", "tabelas_apoio/deflator_ipca_base2024.csv", "Conversão de valores nominais para reais de 2024."],
    ],
    headers=["Dimensão", "Fonte no projeto", "Uso na análise"],
)
add_table(
    doc,
    [
        ["FSA direto", fmt_num((obra["investimento_fsa_deflac"] > 0).sum()), fmt_money(obra["investimento_fsa_deflac"].sum())],
        ["Renúncia fiscal", fmt_num((obra["investimento_renuncia_total_deflac"] > 0).sum()), fmt_money(obra["investimento_renuncia_total_deflac"].sum())],
        ["Bilheteria positiva", fmt_num((obra["bilheteria_deflac"] > 0).sum()), fmt_money(obra["bilheteria_deflac"].sum())],
        ["Outras janelas positivas", fmt_num((obra["outras_janelas_deflac"] > 0).sum()), fmt_money(obra["outras_janelas_deflac"].sum())],
        ["ROI Internacional >= 13", fmt_num((obra["roi_internacional_0_100"] >= 13).sum()), "score composto, não financeiro"],
        ["Crítica com pelo menos 2 fontes", fmt_num((obra["critica_n_fontes"] >= 2).sum()), "índice 1-5"],
    ],
    headers=["Dimensão", "N obras", "Volume/observação"],
)

doc.add_heading("3. Metodologia de cruzamento", level=1)
method_paragraphs = [
    "A chave preferencial para obras é o CPB quando disponível. Quando a fonte externa não traz CPB, a pipeline usa título normalizado: remoção de acentos, padronização de caixa, limpeza de pontuação e comparação por título canônico. Em casos de alteração de título entre projeto e lançamento, a tabela de aliases em tabelas_apoio/titulo_aliases_fsa.csv permite mapear título final para título do projeto FSA.",
    "Projetos FSA são cruzados com obras por título normalizado, CNPJ do requerente quando disponível e chamadas registradas. O resultado alimenta valor FSA nominal, chamada principal, lista de todas as chamadas e valor FSA deflacionado.",
    "Renúncia fiscal é incorporada a partir dos projetos SALIC e bases de fomento indireto. Os valores de Art. 3/3-A/39 e outros mecanismos são agregados por obra; o investimento total público corresponde a FSA deflacionado + renúncia deflacionada.",
    "Produtoras são identificadas por CNPJ e filtradas para agentes independentes brasileiros a partir das bases de produtores e agentes econômicos. A agregação por produtora soma obras, investimentos, receitas, pontuação internacional e atributos de direção.",
    "Bilheteria nacional é usada como principal medida de receita de mercado. Outras janelas entram como estimativas auxiliares quando existentes; por isso, a receita total deve ser lida como aproximação de retorno econômico observável, não como contabilidade completa da obra.",
    "A circulação internacional cruza três famílias de sinais: festivais, admissões europeias Lumière/CNC e presença VOD internacional. Esses sinais são convertidos em um score de 0 a 100 chamado ROI Internacional; apesar do nome, ele mede alcance/circulação, não retorno financeiro.",
    "A fórmula operacional do ROI Internacional é composta por 70 pontos de festivais, 20 pontos de admissões europeias Lumière/CNC e 10 pontos de VOD internacional. O componente de festivais é limitado por cap de 350 pontos antes da normalização; o componente Lumière usa transformação logarítmica log(1 + admissões), normalizada por uma referência de 2,5 milhões de admissões; o componente VOD usa o número de países, limitado a 20 países para saturar os 10 pontos.",
    "O novo entendimento metodológico deste relatório é que o retorno internacional empresarial só é reconhecido quando o ROI Internacional máximo da produtora é maior ou igual a 13. Scores positivos abaixo de 13 permanecem registrados como evidência granular de circulação da obra, mas não qualificam a empresa como Retorno Internacional nem como Duplo Retorno.",
]
for p in method_paragraphs:
    doc.add_paragraph(p)

doc.add_heading("4. Definições analíticas", level=1)
add_table(
    doc,
    [
        ["ROI doméstico FSA", "Receita total estimada / investimento FSA deflacionado."],
        ["ROI doméstico total", "Receita total estimada / investimento público total deflacionado, incluindo FSA e renúncia."],
        ["ROI Internacional", "Score composto 0-100: festivais internacionais, Lumière/CNC e VOD internacional. Não é ROI financeiro."],
        ["Retorno internacional qualificado", "Para produtoras, exige ROI Internacional máximo >= 13 em pelo menos uma obra."],
        ["Duplo Retorno", "Receita total >= R$ 2,5 mi e ROI Internacional máximo >= 13."],
        ["Retorno Doméstico", "Receita total >= R$ 10 mi, ou ROI doméstico > 0,6 com receita >= R$ 2,5 mi, sem ROI Internacional máximo >= 13."],
        ["Retorno Internacional", "ROI Internacional máximo >= 13, sem cumprir critério de Retorno Doméstico."],
        ["Fomento Baixo Retorno", "Investimento total > R$ 5 mi e ausência de retorno doméstico/internacional qualificado."],
        ["Pequeno Porte", "Demais produtoras, em geral com menor investimento total e menor evidência de retorno mensurável."],
    ],
    headers=["Conceito", "Definição operacional"],
)

doc.add_heading("5. Ressalvas e limites de inferência", level=1)
limitations = [
    "A análise é descritiva e associativa. Não estima causalidade dos mecanismos de fomento.",
    "A cobertura de bilheteria, festivais, VOD, crítica e bases internacionais é desigual; ausência de dado não é evidência de ausência de resultado.",
    "O score internacional é uma métrica construída. Ele combina prestígio/circulação, mas não mede venda internacional, margem, lucro ou recuperação financeira externa.",
    "Estimativas de outras janelas dependem das hipóteses e fontes disponíveis na pipeline; devem ser lidas como aproximações.",
    "Cruzamentos por título normalizado podem gerar falsos positivos ou falsos negativos em casos de homônimos, mudanças de título, coproduções e metadados incompletos.",
    "A agregação por produtora pode atribuir resultados a empresas com participação heterogênea em uma mesma obra, pois nem sempre há informação pública completa sobre participação econômica proporcional.",
    "Valores deflacionados reduzem distorções temporais, mas não eliminam mudanças estruturais de mercado entre os anos analisados.",
]
for item in limitations:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6. Visão geral: tamanho do sistema e retorno agregado", level=1)
doc.add_paragraph(
    f"Esta seção apresenta a dimensão do universo analisado, o volume de investimento e a cobertura dos "
    f"principais indicadores. O recorte de retorno usa {fmt_num(len(obra_ret))} obras com bilheteria positiva, "
    f"abrangendo {n_anos_serie} anos de série histórica e {fmt_money(total_inv)} de investimento total deflacionado."
)
add_figure(doc, *chart["01_serie_anual"])
doc.add_paragraph(
    f"A série anual mostra que o pico de investimento deflacionado ocorreu em {peak_inv_year} "
    f"({fmt_money(peak_inv_val)}), enquanto a receita máxima se concentrou em {peak_rec_year}. "
    f"A defasagem de 2 a 3 anos entre investimento em produção e receita de lançamento é estrutural "
    f"no setor audiovisual: obras financiadas em determinado ano chegam ao mercado no ciclo seguinte. "
    f"O ROI doméstico agregado do período é {fmt_ratio(roi_global)} — abaixo de 1x, o que é esperado "
    f"em sistemas de fomento cultural cujo objetivo não se resume ao retorno financeiro. "
    f"A bilheteria é uma das métricas de resultado, não a única, e parte do valor gerado pelo sistema "
    f"é capturado por indicadores de circulação internacional, qualidade crítica e impacto cultural."
)
add_figure(doc, *chart["13_cobertura"])
doc.add_paragraph(
    f"O gráfico de cobertura empírica quantifica, para cada tipo de evidência, quantas obras têm dados "
    f"disponíveis na base consolidada. A bilheteria é o dado mais completo; festivais e Lumière cobrem "
    f"subconjuntos menores, mas são decisivos para medir circulação internacional. "
    f"A crítica com pelo menos duas fontes cobre {fmt_num(len(crit_obras))} obras — parcialidade que deve "
    f"ser considerada ao interpretar os resultados de Soft Power. "
    f"Ausência de dado em qualquer dimensão não é evidência de ausência de resultado: obras sem registro "
    f"de bilheteria podem ter circulado em festivais, e obras sem registro Lumière podem ter tido "
    f"distribuição internacional fora da Europa."
)

doc.add_heading("7. Retorno doméstico e composição do financiamento", level=1)
doc.add_paragraph(
    "Este bloco compara o retorno doméstico entre os quatro grupos de mecanismo de financiamento. "
    "A unidade é a obra, classificada por composição de capital público: FSA puro, renúncia pura "
    "ou arranjos mistos. A comparação mede carteiras observadas — não efeito causal isolado de cada instrumento."
)
add_figure(doc, *chart["02_investimento_grupo"])
doc.add_paragraph(
    f"O grupo '{top_grupo['grupo_mecanismo']}' concentra o maior volume de investimento deflacionado "
    f"({fmt_money(top_grupo['investimento'])}), refletindo tanto o número de obras quanto o porte médio "
    f"dos projetos nessa modalidade. A distribuição entre grupos evidencia que o sistema não é dominado "
    f"por um único instrumento: FSA e renúncia coexistem em volumes comparáveis, e os arranjos mistos "
    f"respondem por parcela relevante do total — indicando que projetos de maior porte mobilizam capital "
    f"de mais de uma fonte pública, seja por exigência regulatória ou estratégia de financiamento."
)
add_figure(doc, *chart["03_roi_grupo"])
doc.add_paragraph(
    f"O '{best_roi_grupo['grupo_mecanismo']}' registra o maior ROI doméstico agregado "
    f"({fmt_ratio(best_roi_grupo['roi'])}), enquanto o '{worst_roi_grupo['grupo_mecanismo']}' tem o "
    f"menor ({fmt_ratio(worst_roi_grupo['roi'])}). Esse padrão reflete diferenças no processo de seleção: "
    f"a renúncia fiscal exige captação de patrocinadores privados, funcionando como filtro de viabilidade "
    f"de mercado antes da produção. O FSA puro, sem esse filtro, financia projetos com menor tração de "
    f"bilheteria mas frequentemente maior ambição artística. "
    f"A análise deve evitar a conclusão de que um mecanismo é 'melhor': cada um serve a objetivos "
    f"distintos de política pública e atinge perfis diferentes de obra e produtora."
)
rows = []
for _, r in group_summary.sort_values("investimento", ascending=False).iterrows():
    rows.append([r["grupo_mecanismo"], fmt_num(r["obras"]), fmt_money(r["investimento"]),
                 fmt_money(r["receita"]), fmt_ratio(r["roi"]), fmt_pct(r["pct_roi_intl_13"])])
add_table(doc, rows, headers=["Grupo", "Obras", "Investimento", "Receita", "ROI agregado", "% ROI intl >=13"])

doc.add_heading("8. Categorias FSA e critério de seleção", level=1)
doc.add_paragraph(
    "As chamadas FSA foram agregadas em categorias analíticas para reduzir a fragmentação histórica dos "
    "editais e permitir comparação longitudinal. A leitura deve separar três eixos independentes: "
    "volume de capital, retorno doméstico e alcance internacional. Uma categoria pode liderar em um eixo "
    "e ser fraca nos outros — essa pluralidade é intencional no desenho do sistema de fomento."
)

if "21_quadrant_categorias" in chart:
    add_figure(doc, *chart["21_quadrant_categorias"])
    if best_cat_roi is not None and best_cat_intl is not None:
        doc.add_paragraph(
            f"O quadrante posiciona cada categoria FSA simultaneamente em retorno doméstico (eixo x) e "
            f"alcance internacional (eixo y); o tamanho do círculo é proporcional ao investimento total. "
            f"'{best_cat_roi['categoria']}' lidera em ROI doméstico ({fmt_ratio(best_cat_roi['roi_agregado'])}), "
            f"enquanto '{best_cat_intl['categoria']}' lidera em alcance internacional "
            f"(score {fmt_num(best_cat_intl['roi_intl'], 1)}). "
            f"Nenhuma categoria ocupa simultaneamente o quadrante superior direito, evidenciando a tensão "
            f"estrutural entre os dois objetivos: categorias com maior retorno doméstico tendem a menor "
            f"internacionalização, e vice-versa. As linhas tracejadas marcam as medianas do conjunto — "
            f"categorias acima e à direita de ambas são as que combinam desempenho acima da mediana nos dois eixos."
        )

add_figure(doc, *chart["05_categoria_investimento"])
if top_cat_inv is not None:
    doc.add_paragraph(
        f"A categoria '{top_cat_inv['categoria']}' concentra o maior volume de investimento total deflacionado "
        f"({fmt_money(top_cat_inv['investimento'])}), com {fmt_num(int(top_cat_inv['obras']))} obras no recorte. "
        f"Essa concentração reflete o peso histórico das chamadas de pontuação por bilheteria, que financiam "
        f"projetos de maior porte com distribuidoras comprometidas com o lançamento comercial. "
        f"As categorias menores em volume — como Automático Festivais e Coprodução Internacional — "
        f"representam instrumentos de nicho com objetivos distintos: reconhecimento artístico e inserção "
        f"em circuitos internacionais, a um custo por projeto muito menor."
    )

add_figure(doc, *chart["06_categoria_roi"])
if best_cat_roi is not None:
    doc.add_paragraph(
        f"Em ROI Total Deflacionado, '{best_cat_roi['categoria']}' apresenta o melhor resultado "
        f"({fmt_ratio(best_cat_roi['roi_agregado'])}), seguida pelas categorias de pontuação por bilheteria "
        f"via distribuidora. O denominador é o capital público total vinculado à obra (FSA + renúncia), "
        f"o que penaliza categorias com forte componente de renúncia. "
        f"A Complementação aparece como eficiente porque financia obras em estágio mais maduro, "
        f"com menor risco de insucesso de mercado. "
        f"Categorias com critério artístico têm ROI doméstico baixo — estruturalmente esperado, pois "
        f"o objetivo dessas linhas é reconhecimento e circulação em festivais, não bilheteria."
    )

add_figure(doc, *chart["07_categoria_intl"])
if best_cat_intl is not None:
    doc.add_paragraph(
        f"No eixo internacional, '{best_cat_intl['categoria']}' lidera com score médio de "
        f"{fmt_num(best_cat_intl['roi_intl'], 1)}, consistente com sua lógica de seleção: "
        f"obras que já receberam reconhecimento em festivais recebem aporte automático, "
        f"criando concentração de prestígio. "
        f"O score total (média × obras) favorece categorias maiores em número de obras, "
        f"enquanto o score médio por obra captura a intensidade de circulação por projeto. "
        f"A análise combinada identifica quais categorias têm poucas obras de altíssimo alcance "
        f"versus muitas obras com alcance moderado — distinção relevante para política de internacionalização."
    )

top_cat_rows = []
for _, r in cat_fsa.sort_values("investimento", ascending=False).head(10).iterrows():
    top_cat_rows.append([r["categoria"], fmt_num(r["obras"]), fmt_money(r["investimento"]),
                         fmt_money(r["receita"]), fmt_ratio(r["roi_agregado"]), fmt_ratio(r["roi_fsa_def"]),
                         fmt_num(r["roi_intl"], 2), fmt_num(r["roi_intl_total"], 2)])
add_table(doc, top_cat_rows, headers=["Categoria", "Obras", "Investimento", "Receita prop.",
                                       "ROI total deflac.", "ROI FSA deflac.", "ROI intl médio", "ROI intl total"])

if "25_critica_categoria" in chart:
    add_figure(doc, *chart["25_critica_categoria"])
    crit_cat_q = (cham.groupby("categoria")
                  .agg(critica=("critica_media", "mean"), obras=("n_obras", "sum"))
                  .reset_index())
    crit_cat_q = crit_cat_q[crit_cat_q["obras"] >= 5]
    if len(crit_cat_q) >= 3:
        top_crit = crit_cat_q.sort_values("critica", ascending=False).iloc[0]
        low_crit = crit_cat_q.sort_values("critica").iloc[0]
        doc.add_paragraph(
            f"O índice crítico médio varia entre categorias: '{top_crit['categoria']}' tem o maior "
            f"({fmt_num(top_crit['critica'], 2)}/5), enquanto '{low_crit['categoria']}' registra o "
            f"menor ({fmt_num(low_crit['critica'], 2)}/5). "
            f"Categorias com critério artístico tendem a avaliações críticas mais altas, pois selecionam "
            f"obras com linguagem elaborada e histórico de reconhecimento em festivais. "
            f"Esse dado complementa a análise de ROI: baixo retorno doméstico combinado com alta "
            f"qualidade crítica indica que a categoria cumpre papel cultural não capturado pela bilheteria."
        )

doc.add_heading("9. Retorno internacional", level=1)
doc.add_paragraph(
    "A análise de retorno internacional mede circulação e reconhecimento externo — não retorno financeiro. "
    "O score é um composto de festivais, admissões Lumière/CNC e VOD internacional. "
    "Na classificação de produtoras, apenas ROI Internacional máximo >= 13 conta como retorno qualificado."
)
add_figure(doc, *chart["04_pct_intl_grupo"])
doc.add_paragraph(
    f"O percentual de obras com ROI Internacional >= 13 é mais alto no grupo com financiamento misto "
    f"majoritariamente FSA, seguido pelo FSA puro. Esse padrão reforça a hipótese de que os editais FSA — "
    f"especialmente os seletivos por critério artístico — favorecem obras com maior circulação internacional. "
    f"A renúncia pura tem o menor percentual de obras com presença qualificada, coerente com seu foco "
    f"em projetos de apelo comercial doméstico. "
    f"No recorte total, {fmt_num(intl_works)} obras atingem score >= 13 e {fmt_num(prod_intl)} produtoras "
    f"têm ao menos uma obra nesse patamar — minoria que concentra a maior parte da representatividade "
    f"internacional do cinema brasileiro em festivais e mercados externos."
)

doc.add_heading("10. Produtoras e clusters", level=1)
doc.add_paragraph(
    "A unidade passa a ser o CNPJ/razão social da produtora, com agregação de obras, investimento, "
    "receita e score internacional. A leitura por cluster separa porte da carteira, retorno doméstico "
    "e evidência internacional qualificada — três dimensões que podem apontar em direções diferentes."
)
add_figure(doc, *chart["08_cluster_counts"])
if biggest_cluster is not None:
    doc.add_paragraph(
        f"O cluster '{biggest_cluster['cluster']}' é o maior em número de produtoras "
        f"({fmt_num(int(biggest_cluster['produtoras']))}), indicando que a maioria das empresas "
        f"opera em escala limitada de bilheteria e sem presença internacional qualificada. "
        f"Os clusters de retorno — Duplo Retorno, Retorno Doméstico e Retorno Internacional — "
        f"concentram minoria das produtoras, mas respondem por parcela desproporcionalmente alta "
        f"do investimento e da receita do sistema. "
        f"Essa pirâmide de concentração é estrutural: poucas empresas reúnem condições de escala, "
        f"recorrência e network para gerar retorno mensurável de forma consistente."
    )

add_figure(doc, *chart["09_cluster_inv_receita"])
doc.add_paragraph(
    "A comparação de investimento e receita por cluster revela o grau de assimetria do sistema. "
    "Duplo Retorno e Retorno Doméstico são os únicos clusters com receita superando o investimento "
    "(ROI > 1x). O Fomento Baixo Retorno concentra alto capital com receita muito abaixo — "
    "o principal ponto de atenção para eficiência alocativa do FSA. "
    "O Pequeno Porte, apesar do maior número de produtoras, tem impacto agregado baixo tanto "
    "em investimento quanto em receita, refletindo a fragilidade estrutural dessas empresas."
)

add_figure(doc, *chart["10_cluster_roi"])
if top_cluster is not None and intl_cluster is not None:
    doc.add_paragraph(
        f"'{top_cluster['cluster']}' lidera em ROI doméstico agregado ({fmt_ratio(top_cluster['roi_agregado'])}), "
        f"enquanto '{intl_cluster['cluster']}' tem o maior score internacional médio "
        f"({fmt_num(intl_cluster['roi_intl_max_med'], 1)}). A linha pontilhada marca ROI = 1x — "
        f"limiar de recuperação do investimento em bilheteria. "
        f"O gráfico confirma que alta internacionalização e alto retorno doméstico tendem a ser "
        f"objetivos concorrentes no sistema atual: as produtoras que maximizam um tendem a sacrificar "
        f"o outro. O Duplo Retorno é a exceção — e é composto por poucas empresas de grande porte."
    )

if "22_scatter_roi_produtoras" in chart:
    add_figure(doc, *chart["22_scatter_roi_produtoras"])
    doc.add_paragraph(
        "O scatter de ROI Doméstico × ROI Internacional por produtora reproduz a visão de matriz de portfolio "
        "do painel interativo. Cada ponto é uma produtora, colorida por cluster; tamanho proporcional "
        "ao número de obras. A concentração de pontos no quadrante inferior esquerdo — baixo retorno "
        "doméstico e baixo alcance internacional — revela que a maioria das produtoras não atinge escala "
        "de retorno em nenhuma das duas dimensões. "
        "As linhas tracejadas (ROI Dom = 1,0x e ROI Intl = 13) delimitam os quadrantes analíticos: "
        "produtoras no quadrante superior direito constituem o Duplo Retorno — poucas, mas com perfil "
        "de excelência bilateral que justifica o interesse como caso de estudo para política pública."
    )

add_figure(doc, *chart["11_scatter_produtoras"])
doc.add_paragraph(
    "O scatter em escala logarítmica de investimento × receita complementa a análise de clusters "
    "mostrando a dispersão interna: produtoras do mesmo cluster podem diferir muito em escala absoluta. "
    "O Duplo Retorno (verde) aparece no canto superior direito — alto investimento e alta receita. "
    "O Pequeno Porte (cinza) forma uma nuvem dispersa no centro-inferior. "
    "A escala logarítmica é necessária porque a distribuição é extremamente assimétrica: "
    "poucas produtoras têm valores 100× maiores que a mediana, tornando escala linear ilegível."
)

rows = []
for _, r in cluster_summary.sort_values("investimento", ascending=False).iterrows():
    rows.append([r["cluster"], fmt_num(r["produtoras"]), fmt_num(r["obras"]),
                 fmt_money(r["investimento"]), fmt_money(r["receita"]),
                 fmt_ratio(r["roi_agregado"]), fmt_num(r["roi_intl_max_med"], 2)])
add_table(doc, rows, headers=["Cluster", "Produtoras", "Obras", "Investimento", "Receita",
                               "ROI agregado", "ROI intl máx médio"])

doc.add_heading("11. Concentração e distribuição do capital", level=1)
doc.add_paragraph(
    "A concentração do FSA é um traço estrutural do sistema. A pergunta muda de retorno para "
    "distribuição: quantas empresas concentram o investimento e qual parcela opera com tickets "
    "muito pequenos para sustentar capacidade institucional. "
    "Concentração não é prova de falha ou eficiência por si só — precisa ser lida à luz do objetivo público."
)
add_figure(doc, *chart["12_lorenz"])
doc.add_paragraph(
    f"A Curva de Lorenz mostra graficamente a desigualdade na distribuição do FSA entre produtoras. "
    f"O Gini calculado no recorte comparável é {fmt_num(gini_all, 3)} para o universo total e "
    f"{fmt_num(gini_pos, 3)} entre produtoras com FSA positivo. "
    f"Para referência, um Gini acima de 0,6 em fomento público indica que a maior parte do capital "
    f"se concentra em grupo muito pequeno de empresas. "
    f"As {fmt_num(n_below_500k)} produtoras com FSA abaixo de R$ 500 mil dificilmente conseguem "
    f"manter estrutura operacional permanente com esse nível de aporte — "
    f"operam de chamada em chamada, sem capacidade de pipeline ou formação de equipe estável. "
    f"As 10 maiores produtoras do recorte concentram {fmt_pct(top10_fsa_share)} do FSA total."
)

if "23_ticket_acumulado" in chart:
    add_figure(doc, *chart["23_ticket_acumulado"])
    pct_500 = n_below_500k / n_pos_fsa * 100 if n_pos_fsa > 0 else 0
    pct_1m  = int(prod_comp["investimento_fsa_deflac"].between(1, 1_000_000).sum()) / n_pos_fsa * 100 if n_pos_fsa > 0 else 0
    doc.add_paragraph(
        f"A curva acumulada evidencia que {fmt_pct(pct_500)} das produtoras com FSA positivo "
        f"acumularam menos de R$ 500 mil no total do período — e {fmt_pct(pct_1m)} abaixo de R$ 1 mi. "
        f"A inflexão em torno de R$ 2 mi delimita onde as produtoras começam a acumular escala recorrente. "
        f"A distância entre R$ 500 mil e R$ 2 mi define uma 'faixa de vulnerabilidade': produtoras "
        f"com operação ativa mas sem estabilidade suficiente para desenvolver pipeline de médio prazo. "
        f"Esse padrão é central no debate sobre fragmentação do setor e sustentabilidade das empresas "
        f"que dependem do FSA como principal fonte de capital de produção."
    )

add_figure(doc, *chart["20_concentracao_tiers"])
doc.add_paragraph(
    "O gráfico de tiers combina a participação percentual no FSA (barras) com o ticket mediano anual "
    "(linha). O Tier E — mais numeroso — concentra o maior número de produtoras com o menor ticket: "
    "a distribuição muito dispersa do capital impede que essas empresas atinjam o patamar mínimo de "
    "operação sustentável. Os tiers superiores (A e B) têm ticket muito acima do necessário e "
    "respondem por parcela relevante dos recursos — o que exige análise de custo-benefício por porte "
    "de projeto e não apenas por volume total de capital alocado."
)

add_table(doc, concentration_rows(prod_comp),
          headers=["Faixa de produtoras com FSA positivo", "N produtoras", "Participação no FSA"])
tier_rows_doc = []
for r in concentration["tier_rows"]:
    tier_rows_doc.append([r["name"], fmt_num(r["n"]), fmt_pct(r["fsa_share"]),
                          f"R$ {fmt_num(r['ticket_med'])} mil",
                          f"R$ {fmt_num(r['fsa_med_obra'], 2)} mi",
                          fmt_pct(r["pct_below_1m"])])
add_table(doc, tier_rows_doc, headers=["Tier", "Produtoras", "% FSA", "Ticket mediano anual",
                                        "FSA mediano/obra", "% abaixo R$1 mi/ano"])
doc.add_paragraph(
    f"No painel de concentração, a base de produtoras independentes com FSA positivo soma "
    f"{fmt_num(concentration['n_prod'])} produtoras, {fmt_money(concentration['total_fsa_m'] * 1e6)} de FSA nominal "
    f"e Gini {fmt_num(concentration['gini'], 3)}. "
    f"{fmt_pct(concentration['pct_below_500k'])} das produtoras operam abaixo de R$ 500 mil/ano e "
    f"{fmt_pct(concentration['pct_below_1m'])} abaixo de R$ 1 milhão/ano."
)

doc.add_heading("12. Diversidade: inscritos e selecionados", level=1)
doc.add_paragraph(
    "Análise de 99 editais FSA (BRDE/FSA 2015–2025) com dados de raça e gênero por função. "
    "24 dos 99 editais tinham política afirmativa (PA). A comparação é descritiva e deve ser lida "
    "como efeito observado nas etapas reportadas — não como medida completa da composição "
    "racial ou de gênero do audiovisual brasileiro."
)

if "24_diversidade_taxas" in chart:
    add_figure(doc, *chart["24_diversidade_taxas"])
    doc.add_paragraph(
        "O gráfico de taxas expõe o efeito das políticas afirmativas com clareza. "
        "Na dimensão racial, editais sem PA mostram taxa de seleção de pessoas negras (26,8%) "
        "abaixo da de brancos (29,9%) — o sistema reproduz o déficit histórico de representação. "
        "Com PA, a relação se inverte: pessoas negras passam a 14,8% ante 9,1% de brancos — "
        "diferença de 63% a favor do grupo historicamente sub-representado. "
        "Na dimensão de gênero, mulheres atingem quase o dobro da taxa dos homens nos editais com PA "
        "(15,0% vs 7,9%). As PAs funcionam — mas cobrem apenas 24% dos editais analisados, "
        "o que significa que o sistema como um todo ainda opera majoritariamente sem correção ativa."
    )

add_figure(doc, *chart["16_genero_direcao"])
doc.add_paragraph(
    "No recorte racial, o percentual de pessoas negras sobe de 15,2% entre inscritos sem PA "
    "para 22,8% com PA e 32,4% entre selecionados com PA. "
    "Esse progressivo aumento ao longo do funil de seleção indica que a PA consegue corrigir "
    "tanto o viés de candidatura quanto o de seleção. "
    "Para análise completa, seria necessário comparar com a composição racial do setor audiovisual — "
    "dado não disponível aqui, mas que contextualizaria se 32,4% representa paridade demográfica "
    "ou ainda sub-representação relativa ao peso do grupo na população."
)

add_figure(doc, *chart["17_diversidade_regional"])
doc.add_paragraph(
    "Na dimensão de gênero, mulheres representam 37,0% das inscritas com PA e 52,6% das selecionadas. "
    "O salto de 37% para 52,6% indica que o critério de seleção com PA favorece ativamente projetos "
    "de direção feminina — não apenas como cota de acesso, mas como critério de mérito ponderado. "
    "O ganho de 15,6 pontos percentuais (pp) com PA supera o ganho racial (+9,6 pp), o que pode "
    "refletir maior clareza das regras de gênero, maior massa de candidatas ou diferenças na "
    "composição dos júris dos editais com e sem PA."
)

doc.add_heading("13. Curtas para longas: trajetória de formação", level=1)
doc.add_paragraph(
    "A análise de curtas em festivais internacionais trata esses registros como sinal de formação "
    "de carreira e entrada em redes de prestígio. O recorte é auxiliar — não cobre todo o universo "
    "de curtas brasileiros, mas mapeia trajetórias de diretores que circularam em festivais de referência "
    "antes ou durante a carreira em longas-metragens."
)
if "27_curtas_uplift" in chart:
    add_figure(doc, *chart["27_curtas_uplift"])
    if curtas_uplift:
        doc.add_paragraph(
            f"Na base geral, {fmt_num(curtas_uplift['base_fest'])} de "
            f"{fmt_num(curtas_uplift['base_total'])} diretores de longas chegam a festival internacional "
            f"({fmt_pct(curtas_uplift['base_pct'])}). Entre diretores com curta selecionado internacionalmente, "
            f"{fmt_num(curtas_uplift['selected_fest'])} de {fmt_num(curtas_uplift['selected_total'])} "
            f"aparecem depois com longa internacional ({fmt_pct(curtas_uplift['selected_pct'])}). "
            f"A diferença é de {fmt_num(curtas_uplift['gain_pp'], 1)} pontos percentuais, "
            f"ou {fmt_num(curtas_uplift['gain_mult'], 1)} vezes a taxa de referência. "
            f"O curta selecionado não prova causalidade isolada, mas funciona como marcador antecipado "
            f"de maior probabilidade de circulação internacional futura."
        )
add_figure(doc, *chart["18_curtas_festivais"])
doc.add_paragraph(
    f"O ranking de festivais mostra a distribuição de {fmt_num(len(curtas))} seleções e premiações "
    f"de curtas entre os festivais mapeados. Cannes, Berlinale e Annecy concentram as seleções "
    f"de maior prestígio; a premiação — distinta da seleção — é mais rara, refletindo o caráter "
    f"competitivo desses festivais de referência mundial. "
    f"A presença em múltiplos festivais de alto prestígio por um mesmo diretor indica circulação "
    f"sustentada em circuito internacional — padrão associado a transições bem-sucedidas para "
    f"longas-metragens com reconhecimento externo. "
    f"A cobertura da base é parcial e tende a capturar os casos mais documentados ou com vínculo "
    f"a projetos FSA: o número real de curtas brasileiros em festivais internacionais é maior."
)

if "26_curtas_por_ano" in chart:
    add_figure(doc, *chart["26_curtas_por_ano"])
    doc.add_paragraph(
        "A série temporal de seleções anuais permite identificar tendências de crescimento ou retração "
        "na presença de curtas brasileiros em festivais internacionais. "
        "Variações anuais podem refletir mudanças na produção nacional, flutuações na seleção dos "
        "festivais ou oscilações na cobertura da base de dados. "
        "O período pós-2020 merece atenção: a pandemia alterou o formato de vários festivais "
        "(edições online e híbridas), o que pode ter reduzido a visibilidade para obras menores "
        "ou, ao contrário, ampliado o acesso a festivais que antes exigiam presença física."
    )

doc.add_heading("Apêndice metodológico: fluxo de dados", level=1)
add_table(
    doc,
    [
        ["1. Ingestão", "Bases ANCINE/FSA, renúncia, bilheteria, festivais, Lumière/VOD, direção e agentes econômicos."],
        ["2. Padronização", "Títulos, CPB, CNPJ, anos, categorias e valores monetários."],
        ["3. Cruzamento", "CPB como chave preferencial; título normalizado e aliases quando CPB não existe."],
        ["4. Deflação", "Valores monetários convertidos para R$ 2024 por IPCA."],
        ["5. Métricas", "Receita total, ROI doméstico, ROI internacional, cobertura e categorias FSA."],
        ["6. Agregações", "Bases por obra, investimento, chamada, produtora e direção."],
        ["7. Classificação", "Produtoras: ROI Internacional máximo >= 13 para retorno internacional qualificado."],
        ["8. Saídas", "Painéis HTML e DOCX renderizados a partir dos datasets consolidados."],
    ],
    headers=["Etapa", "Procedimento"],
)

doc.add_heading("Reprodutibilidade", level=1)
doc.add_paragraph(
    "Este documento foi gerado por scripts/08_enriquecer_doc_analise.py a partir dos datasets em "
    "resultados/datasets/ e da tabela consolidada em resultados/tabela_consolidada_obras.xlsx. "
    "Para reproduzir os números após mudanças de metodologia, rode a pipeline principal e depois este script."
)

try:
    doc.save(OUT_DOC)
    saved_doc = OUT_DOC
except PermissionError:
    saved_doc = OUT_DOC.with_name(f"{OUT_DOC.stem} - corrigido{OUT_DOC.suffix}")
    doc.save(saved_doc)
    print(f"Arquivo principal bloqueado para escrita; salvo como copia corrigida: {saved_doc}")

print(f"Documento enriquecido salvo em: {saved_doc}")
print(f"Gráficos gerados em: {CHART_DIR}")

