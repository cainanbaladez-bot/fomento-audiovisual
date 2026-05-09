from __future__ import annotations

import html
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output_final"

TRECHOS_DOCX = Path(r"C:\Users\INTEL\Desktop\trechos analise de dados.docx")
ANALISE_DOCX = OUTPUT / "Análise de Dados - Fomento Audiovisual Brasileiro.docx"
POLITICA_DOCX = OUTPUT / "Uma política de fomento baseada em evidências_v6.docx"

ANALISE_HTML = OUTPUT / "Análise de Dados - Fomento Audiovisual Brasileiro.html"
POLITICA_HTML = OUTPUT / "Uma política de fomento baseada em evidências_v6.html"

MERGE_MARKER = "18. Trechos complementares incorporados"


POLICY_SECTION_TITLES = {
    "A sinergia FSA e renúncia fiscal no mercado doméstico",
    "O que o retorno das produtoras revela",
    "A estrutura que o fomento (des)construiu",
    "A instabilidade tem endereço: o CGFSA",
    "O que o FSA poderia ser",
}

POLICY_SUBSECTION_TITLES = {
    "De forma geral as distribuidoras selecionam melhor do que as produtoras.",
    "Mas os dados revelam dois perfis completamente distintos:",
    "E as chamadas de ampla concorrência?",
    "A proliferação das produtoras vs sustentabilidade",
    "A Diversidade que temos",
    "Joint Ventures, as SPEs brasileiras",
    "Previsibilidade e planejamento de longo prazo.",
    "Novas possibilidades de investimento em infraestrutura e empregos qualificados.",
    "Renovação depende do desempenho da própria carteira.",
    "Seletivos",
    "SUAT Internacional Curtas",
    "Ampla concorrência a partir do desenvolvimento",
    "Uma distribuição possível:",
}


EVIDENCE_LINKS = [
    {
        "phrases": [
            "renúncia fiscal via artigos 3º e 3ºA",
            "SUAT Major",
            "volumes maiores de P&A",
            "linha de comercialização",
            "regulação do streaming",
        ],
        "target": "7. Retorno doméstico e composição do financiamento",
        "label": "Ver FSA e renúncia",
        "panel_hash": "#retorno-domestico",
    },
    {
        "phrases": [
            "60 chamadas diferentes",
            "60 denominações de chamadas distintas",
        ],
        "target": "2. Descrição dos dados",
        "label": "Ver base e categorias",
        "panel_hash": "#categorias",
    },
    {
        "phrases": [
            "27 obras responsáveis por 75% da renda",
            "50% de toda a renda gerada em salas de cinema",
            "94% da renda de filmes nacionais em salas de cinema com fomento público",
        ],
        "target": "7. Retorno doméstico e composição do financiamento",
        "label": "Ver retorno doméstico",
        "panel_hash": "#retorno-domestico",
    },
    {
        "phrases": [
            "FSA Pontuação Bilheteria com leitura humana de roteiro e distribuidora como proponente",
            "ROI doméstico agregado de 0,94×",
            "ROI doméstico agregado de 1,07×",
            "ROI internacional médio de 3,4",
            "ROI internacional médio de 3,0",
        ],
        "target": "8. Categorias FSA e critério de seleção",
        "label": "Ver categorias FSA",
        "panel_hash": "#categorias",
    },
    {
        "phrases": [
            "filmes selecionados por critério de bilheteria doméstica desempenham melhor",
            "selecionados por critérios de festivais realmente desempenham melhor em internacionalização",
            "Pontuação Festivais e Roteiro",
            "Automático Festivais",
        ],
        "target": "9. Retorno internacional",
        "label": "Ver retorno internacional",
        "panel_hash": "#retorno-internacional",
    },
    {
        "phrases": [
            "ROI internacional 49% superior",
            "ROI internacional 65% superior",
            "3,41 vs. 2,28",
            "3,01 vs. 1,82",
            "distribuidora como proponente",
        ],
        "target": "8. Categorias FSA e critério de seleção",
        "label": "Ver comparativo por proponente",
        "panel_hash": "#categorias",
    },
    {
        "phrases": [
            "mediana fica com ROI Doméstico acima de 1",
            "investimento em comercialização/distribuição",
        ],
        "target": "8. Categorias FSA e critério de seleção",
        "label": "Ver comercialização",
        "panel_hash": "#categorias",
    },
    {
        "phrases": [
            "Duplo Retorno",
            "Retorno Doméstico",
            "Retorno Internacional",
            "Fomento Baixo Retorno",
            "Pequeno Porte",
            "ROI Internacional máximo ≥ 13",
        ],
        "target": "10. Produtoras e clusters",
        "label": "Ver clusters",
        "panel_hash": "#produtoras",
    },
    {
        "phrases": [
            "116 produtoras",
            "154 produtoras",
            "ROI doméstico mediano de 0,11×",
            "ROI internacional mediano zero",
        ],
        "target": "10. Produtoras e clusters",
        "label": "Ver baixo retorno",
        "panel_hash": "#produtoras",
    },
    {
        "phrases": [
            "412 obras com investimento direto do FSA",
            "171 — 41,5%",
            "R$ 295 milhões",
            "capital FSA sem retorno",
        ],
        "target": "11. Concentração e distribuição do capital",
        "label": "Ver capital sem retorno",
        "panel_hash": "#concentracao",
    },
    {
        "phrases": [
            "número de obras produzidas com fomento público cresceu 5,5 vezes",
            "de 72 para 397 por ano",
            "número de produtoras ativas cresceu 5 vezes",
            "investimento mediano por obra caiu 58%",
        ],
        "target": "11. Concentração e distribuição do capital",
        "label": "Ver concentração",
        "panel_hash": "#concentracao",
    },
    {
        "phrases": [
            "coeficiente de Gini",
            "Curva de Lorenz",
            "top 10%",
            "metade das produtoras",
            "54,5% de todo o investimento",
        ],
        "target": "11. Concentração e distribuição do capital",
        "label": "Ver distribuição",
        "panel_hash": "#concentracao",
    },
    {
        "phrases": [
            "Rio de Janeiro e São Paulo juntos concentram 66%",
            "distribuição territorial",
        ],
        "target": "15. Distribuição territorial do investimento",
        "label": "Ver território",
        "panel_hash": "#concentracao",
    },
    {
        "phrases": [
            "3,5× de probabilidade adicional",
            "3,5× mais chance",
            "janela ótima de 3 a 6 anos",
            "curtas em festival tier-1",
            "SUAT Internacional Curtas",
            "curta selecionado para os grandes festivais",
            "cena de curtas no brasil desde 2010",
        ],
        "target": "13. Curtas para longas: trajetória de formação",
        "label": "Ver curtas para longas",
        "panel_hash": "#curtas-longas",
    },
    {
        "phrases": [
            "críticas nacionais e internacionais",
            "citações acadêmicas via OpenAlex",
            "Adoro Cinema",
            "OpenAlex",
        ],
        "target": "14. Soft power e qualidade crítica",
        "label": "Ver crítica e soft power",
        "panel_hash": "#soft-power",
    },
    {
        "phrases": [
            "gênero e raça do BRDE",
            "A Diversidade que temos",
            "política afirmativa",
            "protagonista diverso",
        ],
        "target": "12. Diversidade: inscritos e selecionados",
        "label": "Ver diversidade",
        "panel_hash": "#diversidade",
    },
]


PANEL_ONLY_LINKS = [
    {
        "phrases": ["críticas nacionais e internacionais", "citações acadêmicas via OpenAlex", "Adoro Cinema", "OpenAlex"],
        "href": "Análise do Retorno do Fomento Público ao Audiovisual Brasileiro (FSA - Renúncia Fiscal)_v2.html#soft-power",
        "label": "Ver soft power",
        "panel_hash": "#soft-power",
    },
    {
        "phrases": ["Joint Ventures, as SPEs brasileiras", "SPE, Sociedade de Propósito Específico"],
        "href": "Análise do Retorno do Fomento Público ao Audiovisual Brasileiro (FSA - Renúncia Fiscal)_v2.html#produtoras",
        "label": "Ver produtoras",
        "panel_hash": "#produtoras",
    },
    {
        "phrases": ["CGFSA", "estabilidade e previsibilidade"],
        "href": "Análise do Retorno do Fomento Público ao Audiovisual Brasileiro (FSA - Renúncia Fiscal)_v2.html#visao-geral",
        "label": "Ver visão geral",
        "panel_hash": "#visao-geral",
    },
]


CSS = """
:root {
  color-scheme: light;
  --bg: #f7f4ed;
  --paper: #fffdf8;
  --paper-2: #f0ece2;
  --ink: #191816;
  --muted: #706a5f;
  --faint: #9a9285;
  --rule: #ded7cb;
  --accent: #245d6f;
  --accent-2: #9b5a2e;
  --accent-soft: #e8f1f3;
  --evidence: #7a4a24;
  --shadow: 0 18px 44px rgba(44, 37, 27, .12);
}
* { box-sizing: border-box; }
html { scroll-behavior: smooth; }
body {
  margin: 0;
  background: var(--bg);
  color: var(--ink);
  font: 17px/1.75 Georgia, "Times New Roman", serif;
}
a { color: var(--accent); text-underline-offset: 3px; }
.layout {
  display: grid;
  grid-template-columns: 310px minmax(0, 1fr);
  min-height: 100vh;
}
.sidebar {
  position: sticky;
  top: 0;
  height: 100vh;
  overflow: auto;
  padding: 28px 22px;
  border-right: 1px solid var(--rule);
  background: rgba(255, 253, 248, .84);
  backdrop-filter: blur(14px);
}
.brand {
  font: 700 12px/1.2 Inter, Segoe UI, Arial, sans-serif;
  text-transform: uppercase;
  letter-spacing: .12em;
  color: var(--accent-2);
  margin-bottom: 10px;
}
.doc-title {
  margin: 0 0 22px;
  font: 650 21px/1.2 Inter, Segoe UI, Arial, sans-serif;
  letter-spacing: 0;
}
.search {
  width: 100%;
  height: 38px;
  border: 1px solid var(--rule);
  border-radius: 7px;
  background: #fff;
  padding: 0 11px;
  margin-bottom: 18px;
  color: var(--ink);
  font: 13px Inter, Segoe UI, Arial, sans-serif;
}
.toc {
  display: flex;
  flex-direction: column;
  gap: 3px;
}
.toc a {
  display: block;
  padding: 7px 9px;
  border-radius: 7px;
  color: var(--muted);
  text-decoration: none;
  font: 500 13px/1.35 Inter, Segoe UI, Arial, sans-serif;
}
.toc a.depth-2 { padding-left: 20px; font-size: 12px; color: var(--faint); }
.toc a.depth-3 { padding-left: 32px; font-size: 12px; color: var(--faint); }
.toc a:hover, .toc a.active { background: var(--accent-soft); color: var(--accent); }
.sidebar-footer {
  margin-top: 24px;
  padding-top: 16px;
  border-top: 1px solid var(--rule);
  color: var(--faint);
  font: 12px/1.5 Inter, Segoe UI, Arial, sans-serif;
}
.content-wrap {
  width: min(100%, 1060px);
  margin: 0 auto;
  padding: 54px 46px 90px;
}
article {
  max-width: 820px;
  margin: 0 auto;
  padding: 56px 64px 72px;
  background: var(--paper);
  border: 1px solid rgba(222, 215, 203, .72);
  box-shadow: var(--shadow);
}
h1, h2, h3, h4 {
  font-family: Inter, Segoe UI, Arial, sans-serif;
  line-height: 1.22;
  letter-spacing: 0;
}
h1 {
  margin: 0 0 28px;
  color: var(--accent);
  font-size: 34px;
}
h2 {
  margin: 46px 0 14px;
  padding-top: 16px;
  border-top: 1px solid var(--rule);
  color: var(--ink);
  font-size: 23px;
}
h3 { margin: 30px 0 12px; font-size: 18px; color: var(--accent-2); }
h4 { margin: 22px 0 8px; font-size: 14px; color: var(--muted); text-transform: uppercase; letter-spacing: .08em; }
p { margin: 0 0 17px; }
ul, ol { margin: 0 0 18px 24px; padding: 0; }
li { margin: 6px 0; }
blockquote {
  margin: 24px 0;
  padding: 16px 20px;
  border-left: 3px solid var(--accent);
  background: var(--paper-2);
  color: var(--muted);
}
table {
  width: 100%;
  border-collapse: collapse;
  margin: 24px 0;
  font: 13px/1.45 Inter, Segoe UI, Arial, sans-serif;
}
th, td { padding: 9px 10px; border-bottom: 1px solid var(--rule); text-align: left; vertical-align: top; }
th { color: var(--accent); background: var(--paper-2); font-size: 12px; text-transform: uppercase; letter-spacing: .06em; }
img { max-width: 100%; height: auto; display: block; margin: 24px auto; border-radius: 6px; }
figure { margin: 26px 0; }
figcaption, .chart-caption {
  color: var(--muted);
  font-size: 14px;
  font-style: italic;
  text-align: center;
}
.evidence-link {
  color: inherit;
  text-decoration: none;
  border-bottom: 2px solid rgba(122, 74, 36, .28);
  background: linear-gradient(transparent 62%, rgba(218, 190, 150, .32) 0);
}
.evidence-link:hover { border-bottom-color: rgba(122, 74, 36, .75); }
.evidence-link::after {
  content: " ver análise";
  margin-left: 3px;
  color: var(--evidence);
  font: 600 11px Inter, Segoe UI, Arial, sans-serif;
  white-space: nowrap;
}
.evidence-panel {
  margin: 26px 0 10px;
  padding: 15px 17px;
  border: 1px solid var(--rule);
  border-radius: 8px;
  background: #fbf7ef;
  font: 13px/1.5 Inter, Segoe UI, Arial, sans-serif;
  color: var(--muted);
}
.evidence-panel strong { display: block; margin-bottom: 8px; color: var(--accent-2); }
.evidence-panel a {
  display: inline-block;
  margin: 4px 8px 4px 0;
  padding: 4px 8px;
  border-radius: 999px;
  background: #fff;
  border: 1px solid var(--rule);
  text-decoration: none;
  color: var(--accent);
}
.hide-by-search { display: none !important; }
mark.search-hit { background: #ffe7a8; color: inherit; padding: 0 2px; }
@media (max-width: 920px) {
  .layout { display: block; }
  .sidebar {
    position: relative;
    height: auto;
    border-right: 0;
    border-bottom: 1px solid var(--rule);
  }
  .content-wrap { padding: 24px 14px 54px; }
  article { padding: 34px 22px 46px; }
  h1 { font-size: 28px; }
}
"""


JS = """
const links = [...document.querySelectorAll('.toc a')];
const headings = links.map(a => document.querySelector(a.getAttribute('href'))).filter(Boolean);
const onScroll = () => {
  let current = headings[0];
  for (const h of headings) {
    if (h.getBoundingClientRect().top < 120) current = h;
  }
  links.forEach(a => a.classList.toggle('active', current && a.getAttribute('href') === '#' + current.id));
};
document.addEventListener('scroll', onScroll, { passive: true });
onScroll();

const search = document.querySelector('.search');
if (search) {
  search.addEventListener('input', () => {
    const q = search.value.trim().toLowerCase();
    document.querySelectorAll('article > *').forEach(el => {
      el.classList.toggle('hide-by-search', q && !el.textContent.toLowerCase().includes(q));
    });
  });
}
"""


POLICY_CSS = """
:root {
  --rail: #f3efe5;
  --rail-ink: #26231f;
  --chip: #fff9ee;
}
body {
  background:
    linear-gradient(90deg, rgba(36, 93, 111, .05), transparent 34%),
    var(--bg);
}
.policy-layout {
  width: min(1900px, 100%);
  margin: 0 auto;
  padding: 18px 18px 56px;
}
.policy-topbar {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 18px;
  max-width: 1320px;
  margin: 0 auto 22px;
  color: var(--muted);
  font: 600 12px/1.3 Inter, Segoe UI, Arial, sans-serif;
  text-transform: uppercase;
  letter-spacing: .12em;
}
.policy-topbar a {
  text-transform: none;
  letter-spacing: 0;
  color: var(--accent);
}
.policy-reader {
  display: grid;
  grid-template-columns: minmax(0, 660px) minmax(700px, 980px);
  align-items: start;
  justify-content: center;
  gap: 24px;
}
.policy-reader article {
  max-width: none;
  margin: 0;
  padding: 56px 54px 76px;
}
.policy-reader article > h1:first-child {
  font-size: clamp(34px, 4vw, 52px);
  line-height: 1.05;
  max-width: 12ch;
  margin-bottom: 26px;
}
.policy-reader article > h1:first-child::after {
  content: "Leitura guiada por evidências: clique nos dados destacados para abrir o contexto visual ao lado.";
  display: block;
  max-width: 620px;
  margin-top: 18px;
  color: var(--muted);
  font: 500 15px/1.55 Inter, Segoe UI, Arial, sans-serif;
}
.policy-reader p {
  font-size: 18px;
}
.policy-reader h2 {
  font-size: 26px;
  margin-top: 58px;
}
.policy-reader h3 {
  font-size: 19px;
}
.evidence-link {
  cursor: pointer;
  border-bottom: 2px solid rgba(36, 93, 111, .36);
  background: linear-gradient(transparent 58%, rgba(36, 93, 111, .13) 0);
}
.evidence-link:hover,
.evidence-link.is-active {
  border-bottom-color: var(--accent);
  background: linear-gradient(transparent 52%, rgba(36, 93, 111, .22) 0);
}
.evidence-link::after {
  content: " dado";
  color: var(--accent);
}
.evidence-panel {
  margin: 22px 0 28px;
  background: #fff8ec;
}
.evidence-panel strong {
  font: 700 12px/1.3 Inter, Segoe UI, Arial, sans-serif;
  text-transform: uppercase;
  letter-spacing: .08em;
}
.evidence-panel a {
  cursor: pointer;
}
.evidence-rail {
  position: sticky;
  top: 12px;
  max-height: calc(100vh - 24px);
  overflow: auto;
  padding: 16px;
  border: 1px solid rgba(222, 215, 203, .9);
  border-radius: 8px;
  background: rgba(255, 253, 248, .9);
  box-shadow: var(--shadow);
}
.rail-kicker {
  color: var(--accent-2);
  font: 700 11px/1.2 Inter, Segoe UI, Arial, sans-serif;
  text-transform: uppercase;
  letter-spacing: .12em;
  margin-bottom: 8px;
}
.rail-title {
  margin: 0 0 8px;
  color: var(--rail-ink);
  font: 750 22px/1.18 Inter, Segoe UI, Arial, sans-serif;
}
.rail-copy {
  margin: 0 0 14px;
  color: var(--muted);
  font: 14px/1.55 Inter, Segoe UI, Arial, sans-serif;
}
.rail-metrics {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 8px;
  margin: 12px 0 14px;
}
.metric {
  padding: 10px;
  border: 1px solid var(--rule);
  border-radius: 7px;
  background: var(--chip);
}
.metric b {
  display: block;
  color: var(--accent);
  font: 750 18px/1.1 Inter, Segoe UI, Arial, sans-serif;
}
.metric span {
  display: block;
  margin-top: 4px;
  color: var(--muted);
  font: 11px/1.3 Inter, Segoe UI, Arial, sans-serif;
}
.panel-map {
  display: flex;
  flex-wrap: wrap;
  gap: 7px;
  margin: 12px 0 14px;
}
.panel-map button {
  min-height: 31px;
  border: 1px solid var(--rule);
  border-radius: 7px;
  background: #fff;
  color: var(--accent);
  cursor: pointer;
  font: 750 11px/1.2 Inter, Segoe UI, Arial, sans-serif;
}
.panel-map button:hover,
.panel-map button.is-active {
  border-color: var(--accent);
  background: #eef7f6;
}
.rail-frame-wrap {
  overflow: hidden;
  border: 1px solid #1d2230;
  border-radius: 8px;
  background: #11131f;
}
.rail-frame-label {
  padding: 8px 10px;
  color: #ddd8cc;
  background: #171a2a;
  font: 600 11px/1.3 Inter, Segoe UI, Arial, sans-serif;
}
.rail-frame {
  display: block;
  width: 100%;
  height: min(78vh, 920px);
  min-height: 760px;
  border: 0;
  background: #11131f;
}
.rail-actions {
  display: flex;
  flex-wrap: wrap;
  gap: 8px;
  margin-top: 14px;
}
.rail-actions a {
  display: inline-flex;
  min-height: 34px;
  align-items: center;
  padding: 7px 10px;
  border: 1px solid var(--rule);
  border-radius: 7px;
  background: #fff;
  color: var(--accent);
  text-decoration: none;
  font: 700 12px/1.2 Inter, Segoe UI, Arial, sans-serif;
}
.rail-empty {
  padding: 26px 18px;
  border: 1px dashed var(--rule);
  border-radius: 8px;
  color: var(--muted);
  background: var(--rail);
  font: 14px/1.55 Inter, Segoe UI, Arial, sans-serif;
}
@media (max-width: 1080px) {
  .policy-layout { padding: 18px 14px 52px; }
  .policy-reader { display: block; }
  .policy-reader article { padding: 34px 22px 46px; }
  .evidence-rail {
    position: relative;
    top: auto;
    max-height: none;
    margin-top: 18px;
  }
}
"""


POLICY_JS = """
const PANEL_FILE = 'Análise do Retorno do Fomento Público ao Audiovisual Brasileiro (FSA - Renúncia Fiscal)_v2.html';
const ANALYSIS_FILE = 'Análise de Dados - Fomento Audiovisual Brasileiro.html';

const defaultProfile = {
  title: 'Visão geral do sistema',
  copy: 'Comece pelo panorama: tamanho da base, retorno agregado, mecanismos de financiamento e leitura inicial do desempenho público.',
  panelHash: '#visao-geral',
  analysisHref: ANALYSIS_FILE + '#visão-geral-tamanho-do-sistema-e-retorno-agregado',
  metrics: [['2.990', 'obras na base analisada'], ['R$ 2,38 bi', 'receita doméstica estimada']]
};

const shortcutProfiles = {
  '#visao-geral': defaultProfile,
  '#categorias': {
    title: 'Categorias FSA e critérios',
    copy: 'Compare como bilheteria, festivais, roteiro, distribuidora e produtora como proponente mudam o perfil de retorno dos filmes.',
    panelHash: '#categorias',
    analysisHref: ANALYSIS_FILE + '#categorias-fsa-e-critério-de-seleção',
    metrics: [['0,94x', 'ROI doméstico do padrão distribuidora'], ['3,41', 'ROI internacional médio']]
  },
  '#retorno-domestico': {
    title: 'Retorno doméstico e mecanismos',
    copy: 'Compare FSA, renúncia fiscal, bilheteria, composição do financiamento e retorno doméstico por mecanismo.',
    panelHash: '#retorno-domestico',
    analysisHref: ANALYSIS_FILE + '#retorno-doméstico-e-composição-do-financiamento',
    metrics: [['0,79x', 'ROI doméstico agregado'], ['R$ 2,38 bi', 'receita doméstica estimada']]
  },
  '#retorno-internacional': {
    title: 'Retorno internacional',
    copy: 'Veja festivais, VOD internacional, Lumière/CNC, circulação em salas europeias e alcance internacional das obras.',
    panelHash: '#retorno-internacional',
    analysisHref: ANALYSIS_FILE + '#retorno-internacional',
    metrics: [['535', 'participações em festivais'], ['31', 'países em VOD Europa']]
  },
  '#produtoras': {
    title: 'Produtoras e clusters',
    copy: 'A leitura por produtora separa empresas de retorno doméstico, retorno internacional, duplo retorno, pequeno porte e baixo retorno.',
    panelHash: '#produtoras',
    analysisHref: ANALYSIS_FILE + '#produtoras-e-clusters',
    metrics: [['39', 'produtoras de duplo retorno'], ['51 / 52', 'doméstico / internacional']]
  },
  '#concentracao': {
    title: 'Concentração do capital',
    copy: 'Veja curva de concentração, tiers de produtoras, tickets anuais e o quanto a fragmentação limita capacidade empresarial.',
    panelHash: '#concentracao',
    analysisHref: ANALYSIS_FILE + '#concentração-e-distribuição-do-capital',
    metrics: [['0,634', 'Gini no recorte de FSA positivo'], ['71,0%', 'abaixo de R$ 1 mi/ano']]
  },
  '#curtas-longas': {
    title: 'Curtas para longas',
    copy: 'Acompanhe a ponte entre curta selecionado em festival, maturação profissional e chance posterior de longa com circulação internacional.',
    panelHash: '#curtas-longas',
    analysisHref: ANALYSIS_FILE + '#curtas-para-longas-trajetória-de-formação',
    metrics: [['23,1%', 'curta e longa internacional posterior'], ['2,2x', 'chance relativa vs. base geral']]
  },
  '#diversidade': {
    title: 'Diversidade e seleção',
    copy: 'Compare taxas de seleção por raça e gênero em editais com e sem política afirmativa.',
    panelHash: '#diversidade',
    analysisHref: ANALYSIS_FILE + '#diversidade-inscritos-e-selecionados',
    metrics: [['32,4%', 'pessoas negras selecionadas com PA'], ['52,6%', 'mulheres selecionadas com PA']]
  },
  '#soft-power': {
    title: 'Soft power e crítica',
    copy: 'Contextualize crítica nacional, crítica internacional, festivais, citações acadêmicas e circulação simbólica das obras.',
    panelHash: '#soft-power',
    analysisHref: PANEL_FILE + '#soft-power',
    metrics: [['OpenAlex', 'citações acadêmicas'], ['Crítica', 'nacional e internacional']]
  }
};

const evidenceProfiles = [
  { test: /curtas|curta selecionado|suat internacional curtas|curtas-longas/i, profile: shortcutProfiles['#curtas-longas'] },
  { test: /soft-power|soft power|crítica|criticas|openalex|adoro cinema/i, profile: shortcutProfiles['#soft-power'] },
  { test: /joint ventures|spe|produtoras-e-clusters|clusters|baixo-retorno/i, profile: shortcutProfiles['#produtoras'] },
  { test: /retorno internacional|retorno-internacional|festival|festivais|vod|lumière/i, profile: shortcutProfiles['#retorno-internacional'] },
  { test: /retorno doméstico|retorno-doméstico|bilheteria|renúncia|mecanismo|fsa e renúncia/i, profile: shortcutProfiles['#retorno-domestico'] },
  { test: /visão geral|cgfsa/i, profile: shortcutProfiles['#visao-geral'] },
  {
    test: /categorias|critério|proponente|comercialização|distribuidora/i,
    title: 'Critério de seleção e categorias FSA',
    copy: 'Compare como bilheteria, festivais, roteiro, distribuidora e produtora como proponente mudam o perfil de retorno dos filmes.',
    panelHash: '#categorias',
    metrics: [['0,94x', 'ROI doméstico do padrão distribuidora'], ['3,41', 'ROI internacional médio']]
  },
  {
    test: /retorno-internacional|curtas|festival|festivais/i,
    title: 'Retorno internacional',
    copy: 'Evidências de circulação internacional: festivais, janela de curtas para longas, Lumière/CNC e VOD internacional.',
    panelHash: '#curtas-longas',
    metrics: [['23,1%', 'diretores com curta e longa internacional posterior'], ['2,2x', 'chance relativa vs. base geral']]
  },
  {
    test: /produtoras-e-clusters|clusters|baixo-retorno/i,
    title: 'Produtoras e clusters',
    copy: 'A leitura por produtora separa empresas de retorno doméstico, retorno internacional, duplo retorno, pequeno porte e baixo retorno.',
    panelHash: '#produtoras',
    metrics: [['39', 'produtoras de duplo retorno'], ['51 / 52', 'doméstico / internacional']]
  },
  {
    test: /concentração|distribuição|capital|gini/i,
    title: 'Concentração do capital',
    copy: 'Veja a curva de concentração, tiers de produtoras, tickets anuais e o quanto a fragmentação limita capacidade empresarial.',
    panelHash: '#concentracao',
    metrics: [['0,634', 'Gini no recorte de FSA positivo'], ['71,0%', 'abaixo de R$ 1 mi/ano']]
  },
  {
    test: /diversidade|inscritos|selecionados/i,
    title: 'Diversidade e políticas afirmativas',
    copy: 'Compare taxas de seleção por raça e gênero em editais com e sem política afirmativa.',
    panelHash: '#diversidade',
    metrics: [['32,4%', 'pessoas negras selecionadas com PA'], ['52,6%', 'mulheres selecionadas com PA']]
  },
  {
    test: /retorno-doméstico|bilheteria|renúncia|mecanismo/i,
    title: 'Retorno doméstico e mecanismos',
    copy: 'Use a visão geral do painel para comparar FSA, renúncia fiscal, bilheteria e composição do financiamento.',
    panelHash: '#visao-geral',
    metrics: [['0,79x', 'ROI doméstico agregado'], ['R$ 2,38 bi', 'receita estimada']]
  }
];

function profileFor(link) {
  if (link.dataset.panelHash && shortcutProfiles[link.dataset.panelHash]) {
    return shortcutProfiles[link.dataset.panelHash];
  }
  const key = `${link.href || ''} ${link.title || ''} ${link.textContent || ''}`;
  const matched = evidenceProfiles.find(p => p.test.test(key));
  if (matched?.profile) return matched.profile;
  return matched || {
    title: link.title || 'Evidência relacionada',
    copy: 'Abra o painel ou a seção técnica para ver o dado em contexto.',
    panelHash: '#visao-geral',
    analysisHref: ANALYSIS_FILE,
    metrics: [['Dados', 'contexto visual no painel'], ['Fonte', 'análise técnica']]
  };
}

function renderRail(p, analysisHref, activeHash) {
  const title = document.querySelector('[data-rail-title]');
  const copy = document.querySelector('[data-rail-copy]');
  const metrics = document.querySelector('[data-rail-metrics]');
  const frame = document.querySelector('[data-rail-frame]');
  const analysis = document.querySelector('[data-rail-analysis]');
  const panel = document.querySelector('[data-rail-panel]');
  if (!title || !copy || !metrics || !frame || !analysis || !panel) return;

  title.textContent = p.title;
  copy.textContent = p.copy;
  metrics.innerHTML = p.metrics.map(([value, label]) => `<div class="metric"><b>${value}</b><span>${label}</span></div>`).join('');
  frame.src = encodeURI(PANEL_FILE + p.panelHash);
  analysis.href = analysisHref || p.analysisHref || ANALYSIS_FILE;
  panel.href = encodeURI(PANEL_FILE + p.panelHash);
  document.querySelectorAll('[data-panel-shortcut]').forEach(button => {
    button.classList.toggle('is-active', button.dataset.hash === (activeHash || p.panelHash));
  });
}

function setRail(link) {
  const p = profileFor(link);
  document.querySelectorAll('.evidence-link.is-active').forEach(el => el.classList.remove('is-active'));
  link.classList.add('is-active');
  renderRail(p, link.href, link.dataset.panelHash || p.panelHash);
}

document.querySelectorAll('.evidence-link, .evidence-panel a').forEach(link => {
  link.addEventListener('click', event => {
    event.preventDefault();
    setRail(link);
    if (window.matchMedia('(max-width: 1080px)').matches) {
      document.querySelector('.evidence-rail')?.scrollIntoView({ behavior: 'smooth', block: 'start' });
    }
  });
});

document.querySelectorAll('[data-panel-shortcut]').forEach(button => {
  button.addEventListener('click', () => {
    document.querySelectorAll('.evidence-link.is-active').forEach(el => el.classList.remove('is-active'));
    const profile = shortcutProfiles[button.dataset.hash] || defaultProfile;
    renderRail(profile, profile.analysisHref, button.dataset.hash);
  });
});

renderRail(defaultProfile, defaultProfile.analysisHref, '#visao-geral');
"""


def copy_run_format(src_run, dst_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size


def add_paragraph_copy(dst: Document, src_p) -> None:
    text = src_p.text.strip()
    if not text:
        return
    style_name = src_p.style.name if src_p.style else "Normal"
    if style_name == "SectionHead":
        dst_p = dst.add_paragraph(style="Heading 2")
    elif style_name == "ChartCaption":
        dst_p = dst.add_paragraph(style="Intense Quote" if "Intense Quote" in dst.styles else "Normal")
        dst_p.alignment = src_p.alignment
    elif style_name == "List Paragraph":
        dst_p = dst.add_paragraph(style="List Bullet")
    else:
        dst_p = dst.add_paragraph(style="Normal")

    for run in src_p.runs:
        if not run.text:
            continue
        dst_run = dst_p.add_run(run.text)
        copy_run_format(run, dst_run)
    if not dst_p.runs:
        dst_p.add_run(text)
    if style_name == "ChartCaption":
        for run in dst_p.runs:
            run.italic = True
            run.font.size = Pt(9)


def merge_trechos() -> Path:
    if not TRECHOS_DOCX.exists():
        raise FileNotFoundError(TRECHOS_DOCX)
    if not ANALISE_DOCX.exists():
        raise FileNotFoundError(ANALISE_DOCX)

    current = Document(str(ANALISE_DOCX))
    current_text = "\n".join(p.text for p in current.paragraphs)
    if MERGE_MARKER in current_text:
        return ANALISE_DOCX

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = ANALISE_DOCX.with_name(f"{ANALISE_DOCX.stem}.backup_{stamp}.docx")
    shutil.copy2(ANALISE_DOCX, backup)

    intro = current.add_paragraph()
    intro.add_run().add_break(WD_BREAK.PAGE)
    current.add_heading(MERGE_MARKER, level=1)

    trechos = Document(str(TRECHOS_DOCX))
    for paragraph in trechos.paragraphs:
        add_paragraph_copy(current, paragraph)

    for table in trechos.tables:
        current._body._element.append(deepcopy(table._element))

    current.save(str(ANALISE_DOCX))
    return backup


def slugify(text: str, used: set[str]) -> str:
    value = text.strip().lower()
    value = re.sub(r"^\d+\.\s*", "", value)
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = re.sub(r"[\s_]+", "-", value).strip("-")
    value = value or "secao"
    base = value
    idx = 2
    while value in used:
        value = f"{base}-{idx}"
        idx += 1
    used.add(value)
    return value


def run_pandoc(src: Path) -> str:
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        result = subprocess.run(
            [
                "pandoc",
                str(src),
                "-o",
                str(tmp_path),
                "--standalone",
                "--embed-resources",
                "--wrap=none",
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr)
        return tmp_path.read_text(encoding="utf-8")
    finally:
        tmp_path.unlink(missing_ok=True)


def clean_title(raw_title: str) -> str:
    return raw_title.replace("_", " ").strip()


def soup_body_from_pandoc(markup: str) -> BeautifulSoup:
    soup = BeautifulSoup(markup, "html.parser")
    body = soup.body or soup
    new = BeautifulSoup("<article></article>", "html.parser")
    article = new.article
    for child in list(body.children):
        if isinstance(child, NavigableString) and not child.strip():
            continue
        if isinstance(child, Tag) and child.name in {"script", "style"}:
            continue
        article.append(child.extract())
    return new


def promote_policy_sections(article: Tag) -> None:
    for p in list(article.find_all("p")):
        text = p.get_text(" ", strip=True)
        if text in POLICY_SECTION_TITLES:
            p.name = "h2"
        elif text in POLICY_SUBSECTION_TITLES:
            p.name = "h3"


def normalize_headings(article: Tag) -> None:
    first_h1 = article.find("h1")
    if first_h1 is None:
        first_p = article.find("p")
        if first_p:
            first_p.name = "h1"
    for p in article.find_all("p"):
        text = p.get_text(" ", strip=True)
        if text.startswith("Figura ") or text in {
            "Investimento vs. Retorno por Grupo de financiamento (R$ bilhões, R$2024)",
            "Ranking de ROI Doméstico por categoria de chamada (FSA 2014–2023)",
            "Ranking de ROI Internacional por categoria de chamada (FSA 2014–2023)",
            "Vocação Comercial vs. Alcance Internacional por categoria — Eixo X: ROI doméstico; Eixo Y: ROI Internacional; Tamanho: nº de obras",
            "Quatro clusters de produtoras — ROI Doméstico vs. ROI Internacional (tamanho = investimento FSA)",
        }:
            p["class"] = p.get("class", []) + ["chart-caption"]


def ensure_heading_ids(article: Tag) -> dict[str, str]:
    used: set[str] = set()
    by_text: dict[str, str] = {}
    for h in article.find_all(re.compile("^h[1-4]$")):
        text = h.get_text(" ", strip=True)
        existing = h.get("id")
        if existing and existing not in used:
            used.add(existing)
            ident = existing
        else:
            ident = slugify(text, used)
            h["id"] = ident
        by_text[text] = ident
        simplified = re.sub(r"^\d+\.\s*", "", text)
        by_text.setdefault(simplified, ident)
    return by_text


def target_href(target: str, analysis_ids: dict[str, str]) -> str | None:
    ident = analysis_ids.get(target)
    if not ident:
        for text, h_id in analysis_ids.items():
            if target.lower() in text.lower() or text.lower() in target.lower():
                ident = h_id
                break
    if not ident:
        return None
    return f"{ANALISE_HTML.name}#{ident}"


def link_phrase_in_node(
    soup: BeautifulSoup,
    node: Tag,
    phrase: str,
    href: str,
    label: str,
    panel_hash: str | None = None,
) -> bool:
    for text_node in list(node.find_all(string=True)):
        parent = text_node.parent
        if parent and parent.name in {"a", "script", "style"}:
            continue
        text = str(text_node)
        pos = text.find(phrase)
        if pos < 0:
            continue
        before, matched, after = text[:pos], text[pos : pos + len(phrase)], text[pos + len(phrase) :]
        link = soup.new_tag("a")
        link["href"] = href
        link["class"] = "evidence-link"
        link["title"] = label
        if panel_hash:
            link["data-panel-hash"] = panel_hash
        link.string = matched
        text_node.replace_with(NavigableString(before), link, NavigableString(after))
        return True
    return False


def apply_evidence_links(soup: BeautifulSoup, article: Tag, analysis_ids: dict[str, str]) -> list[tuple[str, str]]:
    applied: list[tuple[str, str]] = []
    seen_targets: dict[str, tuple[str, str | None]] = {}
    for item in EVIDENCE_LINKS:
        href = target_href(item["target"], analysis_ids)
        if not href:
            continue
        for phrase in item["phrases"]:
            for block in article.find_all(["p", "li"]):
                if block.find_parent("a"):
                    continue
                panel_hash = item.get("panel_hash")
                if link_phrase_in_node(soup, block, phrase, href, item["label"], panel_hash):
                    applied.append((item["label"], href))
                    seen_targets[item["label"]] = (href, panel_hash)
                    break

    if seen_targets:
        panel = soup.new_tag("div")
        panel["class"] = "evidence-panel"
        strong = soup.new_tag("strong")
        strong.string = "Atalhos para as análises de dados citadas neste texto"
        panel.append(strong)
        for label, (href, panel_hash) in seen_targets.items():
            a = soup.new_tag("a", href=href)
            if panel_hash:
                a["data-panel-hash"] = panel_hash
            a.string = label
            panel.append(a)
        first_heading = article.find("h1")
        if first_heading:
            first_heading.insert_after(panel)
        else:
            article.insert(0, panel)
    return applied


def apply_panel_only_links(soup: BeautifulSoup, article: Tag) -> list[tuple[str, str]]:
    applied: list[tuple[str, str]] = []
    for item in PANEL_ONLY_LINKS:
        for phrase in item["phrases"]:
            for block in article.find_all(["p", "li", "h2", "h3"]):
                if block.find_parent("a"):
                    continue
                if link_phrase_in_node(soup, block, phrase, item["href"], item["label"], item.get("panel_hash")):
                    applied.append((item["label"], item["href"]))
                    break
    return applied


def build_toc(article: Tag) -> str:
    items = []
    for h in article.find_all(["h1", "h2", "h3"]):
        text = h.get_text(" ", strip=True)
        if not text:
            continue
        depth = {"h1": 1, "h2": 2, "h3": 3}[h.name]
        items.append(
            f'<a class="depth-{depth}" href="#{html.escape(h["id"], quote=True)}">'
            f"{html.escape(text)}</a>"
        )
    return "\n".join(items)


def render_page(article_soup: BeautifulSoup, title: str) -> str:
    article = article_soup.article
    toc = build_toc(article)
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)}</title>
<style>{CSS}</style>
</head>
<body>
<div class="layout">
  <aside class="sidebar">
    <div class="brand">Fomento Audiovisual</div>
    <h1 class="doc-title">{html.escape(title)}</h1>
    <input class="search" type="search" placeholder="Filtrar no documento">
    <nav class="toc" aria-label="Seções do documento">
      {toc}
    </nav>
    <div class="sidebar-footer">Gerado automaticamente a partir do DOCX. Atualize o DOCX e rode <code>python scripts/gerar_docs_html_evidencias.py</code>.</div>
  </aside>
  <main class="content-wrap">
    {str(article)}
  </main>
</div>
<script>{JS}</script>
</body>
</html>
"""


def render_policy_page(article_soup: BeautifulSoup, title: str) -> str:
    article = article_soup.article
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)}</title>
<style>{CSS}</style>
<style>{POLICY_CSS}</style>
</head>
<body>
<div class="policy-layout">
  <header class="policy-topbar">
    <div>Fomento Audiovisual · leitura guiada por dados</div>
    <a href="Análise de Dados - Fomento Audiovisual Brasileiro.html" target="_blank" rel="noopener">Abrir análise técnica completa</a>
  </header>
  <main class="policy-reader">
    {str(article)}
    <aside class="evidence-rail" aria-label="Dados relacionados">
      <div class="rail-kicker">Dados relacionados</div>
      <h2 class="rail-title" data-rail-title>Escolha um dado no texto</h2>
      <p class="rail-copy" data-rail-copy>Clique em um trecho destacado para abrir o contexto visual sem interromper a leitura.</p>
      <div class="rail-metrics" data-rail-metrics>
        <div class="metric"><b>Texto</b><span>argumento principal</span></div>
        <div class="metric"><b>Dados</b><span>evidência ao lado</span></div>
      </div>
      <div class="panel-map" aria-label="Atalhos do painel">
        <button type="button" data-panel-shortcut data-hash="#visao-geral">Visão geral</button>
        <button type="button" data-panel-shortcut data-hash="#retorno-domestico">Retorno doméstico</button>
        <button type="button" data-panel-shortcut data-hash="#retorno-internacional">Retorno internacional</button>
        <button type="button" data-panel-shortcut data-hash="#categorias">Categorias</button>
        <button type="button" data-panel-shortcut data-hash="#produtoras">Produtoras</button>
        <button type="button" data-panel-shortcut data-hash="#concentracao">Concentração</button>
        <button type="button" data-panel-shortcut data-hash="#curtas-longas">Curtas</button>
        <button type="button" data-panel-shortcut data-hash="#diversidade">Diversidade</button>
        <button type="button" data-panel-shortcut data-hash="#soft-power">Soft power</button>
      </div>
      <div class="rail-frame-wrap">
        <div class="rail-frame-label">Mini painel interativo</div>
        <iframe class="rail-frame" data-rail-frame title="Mini painel de dados" loading="lazy"></iframe>
      </div>
      <div class="rail-actions">
        <a data-rail-analysis href="Análise de Dados - Fomento Audiovisual Brasileiro.html" target="_blank" rel="noopener">Abrir seção na análise</a>
        <a data-rail-panel href="Análise do Retorno do Fomento Público ao Audiovisual Brasileiro (FSA - Renúncia Fiscal)_v2.html" target="_blank" rel="noopener">Abrir painel completo</a>
      </div>
    </aside>
  </main>
</div>
<script>{POLICY_JS}</script>
</body>
</html>
"""


def convert_docx(src: Path, dst: Path, kind: str, analysis_ids: dict[str, str] | None = None) -> dict[str, str]:
    markup = run_pandoc(src)
    article_soup = soup_body_from_pandoc(markup)
    article = article_soup.article
    if kind == "policy":
        promote_policy_sections(article)
    normalize_headings(article)
    ids = ensure_heading_ids(article)
    if kind == "policy" and analysis_ids:
        apply_evidence_links(article_soup, article, analysis_ids)
        apply_panel_only_links(article_soup, article)
    title = clean_title(src.stem)
    renderer = render_policy_page if kind == "policy" else render_page
    dst.write_text(renderer(article_soup, title), encoding="utf-8")
    return ids


def main() -> None:
    merge_result = merge_trechos()
    if merge_result == ANALISE_DOCX:
        print(f"[OK] Trechos já estavam incorporados em: {ANALISE_DOCX}")
    else:
        print(f"[OK] Backup criado: {merge_result}")
        print(f"[OK] Trechos incorporados em: {ANALISE_DOCX}")

    analysis_ids = convert_docx(ANALISE_DOCX, ANALISE_HTML, "analysis")

    print(f"[OK] HTML gerado: {ANALISE_HTML}")
    print(f"[INFO] Politica v6.html nao e gerado aqui — use atualizar_texto.py")


if __name__ == "__main__":
    main()
